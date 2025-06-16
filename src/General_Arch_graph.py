# =============================================================================
# 1. Imports
# =============================================================================
import os
import uuid
import json
import base64
import shutil # NEW IMPORT for file copying
import cv2 # 新增或確保 cv2 在頂部被引入
from typing import Dict, List, Any, Annotated, Literal, Union, Optional, Tuple
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv
from typing_extensions import TypedDict
from contextlib import asynccontextmanager
import traceback # Added for error printing
from jinja2 import Environment, FileSystemLoader # NEW IMPORT
from docx import Document as DocxDocument # NEW IMPORT for Word document creation
from docx.shared import Inches, Pt # NEW IMPORT for sizing
from docx.enum.text import WD_ALIGN_PARAGRAPH # NEW IMPORT for alignment
import plotly.graph_objects as go # NEW IMPORT for Sankey diagram
from PIL import Image # Ensure PIL is imported at the top level

# LangChain/LangGraph Imports
from langchain_core.messages import AIMessage, HumanMessage, BaseMessage
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough, RunnableLambda, RunnableConfig
from langchain_core.documents import Document
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_anthropic import ChatAnthropic
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.memory import ConversationBufferWindowMemory, VectorStoreRetrieverMemory
from langchain_community.vectorstores import Chroma
from langgraph.graph import StateGraph, END
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder, PromptTemplate # Import necessary prompt classes
from langchain_core.output_parsers import JsonOutputParser # NEW IMPORT
import re # ADD THIS

# Custom Imports
# --- Core Tools ---
from src.tools.ARCH_rag_tool import ARCH_rag_tool
from src.tools.img_recognition import img_recognition
from src.tools.video_recognition import video_recognition
from src.tools.gemini_image_generation_tool import generate_gemini_image
from src.tools.gemini_search_tool import perform_grounded_search
# --- ComfyUI Tools ---
from src.tools.model_render_image import model_render_image # Ensure this is the one being imported
from src.tools.generate_3D import generate_3D
# --- MODIFIED: Remove simulate_future_image import ---
# from src.tools.simulate_future_image import simulate_future_image

# --- MODIFIED ---
# Import ConfigSchema, but remove direct dependency on FullConfig, AgentConfig etc. for runtime
from src.configuration import ConfigManager, ModelConfig, MemoryConfig, initialize_llm, ConfigSchema
# --- <<< 新增：從 state 導入 >>> ---
from src.state import WorkflowState, TaskState
# --- <<< 結束新增 >>> ---

# --- <<< 新增：從子圖檔案導入 >>> ---
from src.AssignTeamsSubgraph import assign_teams
from src.EvaluationSubgraph import evaluation_teams 
# --- <<< 結束新增 >>> ---

# --- <<< 新增：從 configuration 導入SankeyStructureAgent相關配置 (雖然目前是直接訪問) >>> ---
# No specific import needed yet as we will access config via _base_default_config_obj and runtime_config

# =============================================================================
# 2. 環境變數與組態載入 (Load static parts if needed)
# =============================================================================
load_dotenv()

# --- MODIFIED ---
# ConfigManager might still be used for *static* config parts like prompts
# or initial non-configurable settings.
# If prompts are static, load them here. Let's assume they are for now.
config_manager = ConfigManager("config.json")
_full_static_config = config_manager.get_full_config() # Load once for static parts
workflow_config_static = _full_static_config.workflow
# --- Get static prompts (example) ---
_pm_prompts = _full_static_config.agents.get("process_management", {}).prompts or {}
_aa_prompts = _full_static_config.agents.get("assign_agent", {}).prompts or {}
_ta_prompts = _full_static_config.agents.get("tool_agent", {}).prompts or {}
_ea_prompts = _full_static_config.agents.get("eva_agent", {}).prompts or {}
_qa_prompts = _full_static_config.agents.get("qa_agent", {}).prompts or {}
# Get static tool descriptions for AssignAgent
_aa_tool_descriptions = _full_static_config.agents.get("assign_agent", {}).parameters.get("tool_descriptions", {})
# Get static tool default parameters for ToolAgent
_ta_tool_configs_static = _full_static_config.agents.get("tool_agent", {}).tools or {}
_ta_file_handling_static = _full_static_config.agents.get("tool_agent", {}).parameters.get("file_handling", {})
# --- END MODIFIED ---

# =============================================================================
# 3. 常數設定
# =============================================================================
OUTPUT_DIR = workflow_config_static.output_directory
RENDER_CACHE_DIR = os.path.join(OUTPUT_DIR, "render_cache")
MODEL_CACHE_DIR = os.path.join(OUTPUT_DIR, "model_cache")
os.makedirs(RENDER_CACHE_DIR, exist_ok=True)
os.makedirs(MODEL_CACHE_DIR, exist_ok=True)
MEMORY_DIR = os.path.join("knowledge", "memory")
os.makedirs(MEMORY_DIR, exist_ok=True)
LLM_OUTPUT_LANGUAGE_DEFAULT = workflow_config_static.llm_output_language

# =============================================================================
# 4. LLM 與記憶體初始化
# =============================================================================
# --- MODIFIED: Use MEMORY_DIR and _full_static_config from common_definitions ---
ltm_embedding_config = _full_static_config.memory.long_term_memory
if ltm_embedding_config.provider == "openai":
     embeddings = OpenAIEmbeddings(model=ltm_embedding_config.model_name)
else:
     print(f"Warning: Unsupported LTM embedding provider '{ltm_embedding_config.provider}'. Defaulting to OpenAI.")
     embeddings = OpenAIEmbeddings(model="text-embedding-3-small")

vectorstore = Chroma(
    persist_directory=MEMORY_DIR, # Use imported MEMORY_DIR
    embedding_function=embeddings,
    collection_name="workflow_memory"
)
ltm_memory_key = "ltm_context"
ltm_input_key = "input"

print(f"Long-term memory vector store initialized. Storing in: {MEMORY_DIR}")
print(f"LTM Retriever 'k' value will be determined at runtime from configuration.")

# =============================================================================
# 5. 狀態定義
# =============================================================================
# --- REMOVED: Definitions moved to src/state.py ---

# =============================================================================
# 6. Agent 類別定義 (PM, Eva, QA)
# =============================================================================
class ProcessManagement:
    # --- MODIFIED __init__ ---
    def __init__(self, long_term_memory_vectorstore=vectorstore): # Assuming 'vectorstore' is globally available or passed
        # Remove LLM and config args if not needed directly for saving
        # self.llm = llm
        # self.config = config
        # self.prompts = config.get("prompts", {}) # Keep prompts if failure analysis still uses them
        # self.max_retries = config.get("parameters", {}).get("max_retries", 3)
        # self.llm_output_language = global_config.get("llm_output_language", "繁體中文")

        # --- Store the vectorstore ---
        if long_term_memory_vectorstore is None:
            raise ValueError("ProcessManagement requires a valid vectorstore instance.")
        self.vectorstore = long_term_memory_vectorstore
        print("ProcessManagement initialized with vectorstore.")

        # Load prompts if needed for failure analysis/interrupt commands
        # You might load the config manager here if necessary
        config_manager = ConfigManager() # Or get instance if singleton
        pm_config = config_manager.get_agent_config("process_management")
        self.prompts = {name: cfg.template for name, cfg in pm_config.prompts.items()} if pm_config else {}
        self.llm_output_language = config_manager.get_workflow_config().llm_output_language

        # Get LLM for potential failure analysis / non-QA interrupts
        self.llm = initialize_llm(pm_config.llm.model_dump()) if pm_config else None # Initialize LLM if needed

    def _create_task_summary(self, task: TaskState) -> str:
        """Creates a textual summary of a task for LTM."""
        # ...(Existing summary logic - seems fine)...
        summary_lines = [
            f"Task ID: {task.get('task_id', 'N/A')}",
            f"Status: {task.get('status', 'N/A')}",
            f"Objective: {task.get('task_objective', 'N/A')}",
            f"Description: {task.get('description', 'N/A')}",
            f"Agent: {task.get('selected_agent', 'N/A')}",
            f"Requires Eval: {task.get('requires_evaluation', 'N/A')}",
            f"Retry Count: {task.get('retry_count', 'N/A')}",
        ]
        # Include summarized inputs if they exist and are simple enough
        if task.get("task_inputs"):
             try:
                  # Limit size/depth of inputs in summary
                  inputs_str = json.dumps(task["task_inputs"], ensure_ascii=False, default=lambda o: "<non-serializable>", indent=None)[:200]
                  summary_lines.append(f"Inputs Summary: {inputs_str}")
             except Exception:
                  summary_lines.append("Inputs Summary: <error serializing>")

        # Include summarized outputs if they exist
        if task.get("outputs"):
             try:
                  outputs_str = json.dumps(task["outputs"], ensure_ascii=False, default=lambda o: "<non-serializable>", indent=None)[:300]
                  summary_lines.append(f"Outputs Summary: {outputs_str}")
             except Exception:
                  summary_lines.append("Outputs Summary: <error serializing>")

        # Include file info
        if task.get("output_files"):
            file_info = [f"{f.get('filename', 'N/A')} ({f.get('type', 'N/A')})" for f in task["output_files"]]
            summary_lines.append(f"Output Files: {', '.join(file_info)}")

        # Include evaluation summary if present
        if task.get("evaluation") and task["evaluation"].get("assessment"):
            summary_lines.append(f"Evaluation: {task['evaluation']['assessment']} - {task['evaluation'].get('feedback', '')[:100]}...")

        # Include error log if present
        if task.get("error_log"):
            summary_lines.append(f"Error Log: {task['error_log'][:150]}...") # Truncate error log

        # Include feedback log if present
        if task.get("feedback_log"):
             summary_lines.append(f"Feedback Log: {task['feedback_log'][:150]}...") # Truncate feedback log

        return "\n".join(summary_lines)

    # --- Method to save task summary to LTM ---
    async def _save_task_to_ltm(self, task: TaskState):
        """Generates a summary and saves a task state to the vectorstore."""
        task_id = task.get('task_id', 'N/A')
        status = task.get('status', 'N/A')
        agent_name = task.get('selected_agent', 'N/A') # Get agent name for logging

        if task_id == 'N/A':
            print("PM LTM Warning: Cannot save task without ID.")
            return

        print(f"PM LTM: Preparing to save task {task_id} (Agent: {agent_name}, Status: '{status}')...")

        # --- <<< ADDED LOGGING for outputs >>> ---
        task_outputs_to_log = task.get('outputs')
        if task_outputs_to_log:
            try:
                # Log a preview, be careful with large outputs
                outputs_preview = json.dumps(task_outputs_to_log, ensure_ascii=False, default=str)[:500]
                print(f"  - Task {task_id} Outputs being saved to LTM (Preview): {outputs_preview}")
                # Specifically check Pinterest output structure if agent matches
                if agent_name == "PinterestMCPCoordinator":
                     pinterest_paths = task_outputs_to_log.get("saved_image_paths")
                     if isinstance(pinterest_paths, list):
                         print(f"    - Pinterest paths found in outputs: Count={len(pinterest_paths)}, First few: {pinterest_paths[:3]}")
                     else:
                         print(f"    - Pinterest paths ('saved_image_paths') key not found or not a list in outputs.")
            except Exception as log_e:
                print(f"  - Warning: Could not serialize/preview outputs for LTM logging: {log_e}")
        else:
            print(f"  - Task {task_id} has no 'outputs' field to save to LTM.")
        # --- <<< END ADDED LOGGING >>> ---

        summary = self._create_task_summary(task)
        # print(f"PM LTM Summary for {task_id}:\n{summary}\n--------------------") # Optional: Debug print summary

        # Create LangChain Document object
        # Metadata helps retrieval filtering later if needed
        metadata = {
            "task_id": task_id,
            "status": status,
            "agent": agent_name,
            "requires_evaluation": task.get("requires_evaluation", False),
            "timestamp": datetime.now().isoformat() # Add timestamp
        }
        doc = Document(page_content=summary, metadata=metadata)

        try:
            await self.vectorstore.aadd_documents([doc]) # Use async add
            print(f"PM LTM: Successfully saved task {task_id} (Status: {status}) to LTM.")
        except Exception as e:
            print(f"PM LTM Error: Failed to save task {task_id} to vectorstore: {e}")

    # --- MODIFIED run method ---
    async def run(self, state: WorkflowState, config: RunnableConfig) -> WorkflowState:
        print("\n=== Process Management Running ===")
        output_state = state.copy()
        output_state["interrupt_result"] = None # Ensure cleared initially

        # ... (打印 initial state 的日誌 - 不變) ...
        initial_tasks = state.get("tasks", [])
        initial_idx = state.get("current_task_index", 0)
        print(f"  PM Start: Received initial current_task_index: {initial_idx}")
        print(f"  PM Start: Received initial task count: {len(initial_tasks)}")
        if 0 <= initial_idx < len(initial_tasks):
             print(f"  PM Start: Initial task at index {initial_idx} ID: '{initial_tasks[initial_idx].get('task_id')}', Status: '{initial_tasks[initial_idx].get('status')}'")
        else:
             print(f"  PM Start: Initial index {initial_idx} is out of bounds or tasks empty.") # MODIFIED for clarity


        tasks = state.get("tasks", [])
        current_idx = state.get("current_task_index", 0)
        interrupt_input = state.get("interrupt_input") # Use original state's input

        # --- 1. Interrupt Processing ---
        if interrupt_input:
            print(f"PM: Interrupt detected: '{interrupt_input[:100]}...'. Invoking LLM for analysis.")
            # ... (LLM interrupt analysis logic - remains the same) ...
            summarized_tasks = [{"task_id": t.get("task_id", "N/A"), "objective": t.get("task_objective", "N/A")} for t in tasks]
            tasks_json_for_interrupt = json.dumps(summarized_tasks, ensure_ascii=False)
            current_task_for_interrupt = tasks[current_idx] if 0 <= current_idx < len(tasks) else None
            
            # --- <<< 修改：過濾當前任務以生成 JSON >>> ---
            current_task_filtered = {}
            if current_task_for_interrupt:
                current_task_filtered = current_task_for_interrupt.copy()
                # 從 task['outputs'] 中移除不需要的鍵
                if "outputs" in current_task_filtered and isinstance(current_task_filtered["outputs"], dict):
                    current_task_outputs_filtered = current_task_filtered["outputs"].copy()
                    current_task_outputs_filtered.pop("mcp_internal_messages", None)
                    current_task_outputs_filtered.pop("grounding_sources", None)
                    current_task_outputs_filtered.pop("search_suggestions", None) 
                    current_task_filtered["outputs"] = current_task_outputs_filtered 
                
                # --- NEW: Filter base64_data from output_files for interrupt prompt ---
                if "output_files" in current_task_filtered and isinstance(current_task_filtered["output_files"], list):
                    filtered_output_files = []
                    for file_dict in current_task_filtered["output_files"]:
                        if isinstance(file_dict, dict):
                            file_copy = file_dict.copy()
                            file_copy.pop("base64_data", None)
                            file_copy.pop("fileName", None)
                            filtered_output_files.append(file_copy)
                        else:
                            filtered_output_files.append(file_dict) # Keep non-dict items as is
                    current_task_filtered["output_files"] = filtered_output_files
                # --- END NEW ---

                # 可以考慮也過濾 task_inputs 如果需要
                # if "task_inputs" in current_task_filtered and isinstance(current_task_filtered["task_inputs"], dict): ...

            current_task_json_for_interrupt = json.dumps(current_task_filtered, ensure_ascii=False, default=lambda o: '<not serializable>') if current_task_filtered else "{}"
            # --- <<< 結束修改 >>> ---

            interrupt_prompt_input = {
                "user_input": state.get("user_input", ""), "interrupt_input": interrupt_input,
                "current_task_index": current_idx, "tasks_json": tasks_json_for_interrupt,
                "current_task_json": current_task_json_for_interrupt, "llm_output_language": self.llm_output_language,
            }
            try:
                if not self.prompts.get("process_interrupt"): raise ValueError("process_interrupt prompt missing")
                if not self.llm: raise ValueError("LLM for PM not initialized.")

                interrupt_chain = self.llm | StrOutputParser()
                interrupt_content = await interrupt_chain.ainvoke(self.prompts["process_interrupt"].format(**interrupt_prompt_input))

                if interrupt_content.startswith("```json"): interrupt_content = interrupt_content[7:-3].strip()
                elif interrupt_content.startswith("```"): interrupt_content = interrupt_content[3:-3].strip()
                interrupt_result = json.loads(interrupt_content)
                action = interrupt_result.get("action", "PROCEED")
                print(f"PM: Parsed LLM interrupt analysis action: {action}, Result: {interrupt_result}") # Log parsed result

                # --- <<< 修改：處理 Command Actions 並控制返回邏輯 >>> ---
                processed_interrupt = False # Flag to track if interrupt caused an immediate return

                if action == "CONVERSATION":
                    print("PM: LLM analysis resulted in CONVERSATION. Setting phase to 'qa'.")
                    output_state["current_phase"] = "qa"
                    output_state["interrupt_result"] = interrupt_result # Keep result for QA entry check
                    # Keep interrupt_input for QA Agent? No, QA Agent reads from context. Clear it.
                    output_state["interrupt_input"] = None
                    print(f"PM (CONVERSATION): Returning state for QA. Index remains {current_idx}.")
                    processed_interrupt = True
                    return output_state # CONVERSATION requires immediate return to route to QA

                elif action == "REPLACE_TASKS":
                    print("PM: Processing REPLACE_TASKS command.")
                    new_tasks_data = interrupt_result.get("new_tasks_list")
                    if isinstance(new_tasks_data, list):
                        created_tasks = []
                        for i, task_data in enumerate(new_tasks_data):
                            if isinstance(task_data, dict):
                                task_id = str(uuid.uuid4())
                                task = TaskState(
                                    task_id=task_id, status="pending",
                                    task_objective=task_data.get("task_objective", f"Replaced Objective {i+1}"),
                                    description=task_data.get("description", f"Replaced Task {i+1}"),
                                    selected_agent=task_data.get("selected_agent", "LLMTaskAgent"),
                                    task_inputs=task_data.get("inputs", {}),
                                    outputs={}, output_files=[], evaluation={}, error_log=None, feedback_log=None,
                                    requires_evaluation=task_data.get("requires_evaluation", False), retry_count=0
                                )
                                if not task["selected_agent"]:
                                    print(f"PM Warning (REPLACE): Task {i} lacks 'selected_agent'. Skipping.")
                                    continue
                                created_tasks.append(task)
                            else: print(f"PM Warning (REPLACE): Invalid item in new_tasks_list at index {i}.")

                        if created_tasks:
                            original_tasks = state.get("tasks", [])
                            interrupt_idx = state.get("current_task_index", 0)
                            completed_prefix = [t for idx, t in enumerate(original_tasks) if idx < interrupt_idx and t.get("status") == "completed"]
                            print(f"PM (REPLACE): Preserving {len(completed_prefix)} completed tasks before index {interrupt_idx}.")
                            output_state["tasks"] = completed_prefix + created_tasks
                            output_state["current_task_index"] = len(completed_prefix)
                            output_state["current_task"] = created_tasks[0].copy() if created_tasks else None
                            print(f"PM (REPLACE): State updated. New index: {output_state['current_task_index']}. New task count: {len(output_state['tasks'])}.")
                        else:
                            print("PM Warning (REPLACE): No valid tasks created. Proceeding.")
                            # No need to explicitly set action = "PROCEED", will fall through
                    else:
                        print("PM Warning (REPLACE): 'new_tasks_list' missing or invalid. Proceeding.")
                        # No need to explicitly set action = "PROCEED", will fall through

                    # --- Common logic for REPLACE/INSERT ---
                    output_state["interrupt_input"] = None # Clear after processing
                    output_state["interrupt_result"] = None
                    processed_interrupt = True
                    print(f"PM (REPLACE): Returning updated state to router after replacing tasks.")
                    return output_state # REPLACE requires immediate return

                elif action == "INSERT_TASKS":
                    print("PM: Processing INSERT_TASKS command.")
                    insert_tasks_data = interrupt_result.get("insert_tasks_list")
                    if isinstance(insert_tasks_data, list):
                        inserted_tasks = []
                        for i, task_data in enumerate(insert_tasks_data):
                            if isinstance(task_data, dict):
                                task_id = str(uuid.uuid4())
                                task = TaskState(
                                    task_id=task_id, status="pending", # Ensure status is pending
                                    task_objective=task_data.get("task_objective", f"Inserted Objective {i+1}"),
                                    description=task_data.get("description", f"Inserted Task {i+1}"),
                                    selected_agent=task_data.get("selected_agent", "LLMTaskAgent"),
                                    task_inputs=task_data.get("inputs", {}),
                                    outputs={}, output_files=[], evaluation={}, error_log=None, feedback_log=None,
                                    requires_evaluation=task_data.get("requires_evaluation", False), retry_count=0
                                )
                                if not task["selected_agent"]:
                                    print(f"PM Warning (INSERT): Task {i} lacks 'selected_agent'. Skipping.")
                                    continue
                                inserted_tasks.append(task)
                                print(f"  - Created inserted task {i+1} (ID: {task_id})")
                            else: print(f"PM Warning (INSERT): Invalid item in insert_tasks_list at index {i}.")

                        if inserted_tasks:
                            original_tasks_before_insert = state.get("tasks", [])
                            insert_after_idx = state.get("current_task_index", 0)
                            valid_insert_point = insert_after_idx + 1
                            print(f"PM (INSERT): Original index before insert: {insert_after_idx}")
                            print(f"PM (INSERT): Inserting {len(inserted_tasks)} tasks at index {valid_insert_point}.")

                            current_tasks_in_output_state = output_state.get("tasks", original_tasks_before_insert)
                            new_sequence = current_tasks_in_output_state[:valid_insert_point] + inserted_tasks + current_tasks_in_output_state[valid_insert_point:]

                            output_state["tasks"] = new_sequence
                            output_state["current_task_index"] = valid_insert_point
                            output_state["current_task"] = inserted_tasks[0].copy()
                            print(f"PM (INSERT): State updated. New index: {output_state['current_task_index']}. New task count: {len(new_sequence)}.")
                            print(f"PM (INSERT): Current task set to ID: {output_state['current_task']['task_id']}")
                        else:
                            print("PM Warning (INSERT): No valid tasks created from LLM response. Proceeding.")
                            # No need to explicitly set action = "PROCEED", will fall through
                    else:
                        print("PM Warning (INSERT): 'insert_tasks_list' missing or invalid in LLM response. Proceeding.")
                        # No need to explicitly set action = "PROCEED", will fall through

                    # --- Common logic for REPLACE/INSERT ---
                    output_state["interrupt_input"] = None # Clear after processing
                    output_state["interrupt_result"] = None
                    processed_interrupt = True
                    print(f"PM (INSERT): Returning updated state to router after inserting tasks.")
                    return output_state # INSERT requires immediate return

                # --- Handle PROCEED (or fallback): Only clear flags, DO NOT return ---
                elif action == "PROCEED":
                    print("PM: PROCEED command received or fallback. Clearing interrupt flags and continuing.")
                    output_state["interrupt_input"] = None
                    output_state["interrupt_result"] = None
                    # --- DO NOT RETURN HERE ---
                    processed_interrupt = False # Indicate we should continue to task processing loop

                # --- Cleanup after handling non-returning actions (like PROCEED) ---
                # This part is now only reached if action was PROCEED
                if not processed_interrupt:
                    print(f"PM: Processed non-returning interrupt command '{action}'. Proceeding to task check loop.")
                    # Flags already cleared above for PROCEED case.

                # --- Catch potential errors during interrupt processing ---
                # Removed the explicit `except Exception as e:` block that was wrapping the action handling,
                # as individual actions returning should handle their own errors or the main function try-except will catch it.
                # Ensure the try-except around the LLM call itself remains.

            except Exception as e:
                print(f"PM Error processing interrupt command: {e}")
                traceback.print_exc()
                output_state["interrupt_input"] = None # Clear on error too
                output_state["interrupt_result"] = None
                # Consider returning state here too, or let it flow to task processing?
                # Let's allow it to flow to task processing for now, which might handle failure.
                print(f"PM: Error during interrupt processing. Allowing flow to task status check.")


        # --- 2. Task Processing Logic ---
        # This part is now reached if:
        # a) There was no interrupt_input initially.
        # b) The interrupt_input resulted in a 'PROCEED' action.
        tasks = output_state.get("tasks", []) # Use potentially updated tasks (though PROCEED doesn't modify them)
        current_idx = output_state.get("current_task_index", 0) # Use potentially updated index (though PROCEED doesn't modify it)

        # --- Initial Workflow Creation (if tasks is empty and no interrupt happened) ---
        if not tasks and not state.get("interrupt_input"): # Only create if no interrupt was processed
            # ... (Existing workflow creation logic) ...
            print("PM: Task list is empty. Generating initial workflow...")
            try:
                if not self.prompts.get("create_workflow"): raise ValueError("create_workflow prompt missing")
                if not self.llm: raise ValueError("LLM for PM not initialized.") # Check LLM
                create_chain = self.llm | StrOutputParser()
                create_input = {"user_input": state.get("user_input", ""), "llm_output_language": self.llm_output_language}
                plan_content = await create_chain.ainvoke(self.prompts["create_workflow"].format(**create_input))

                if plan_content.startswith("```json"): plan_content = plan_content[7:-3].strip()
                elif plan_content.startswith("```"): plan_content = plan_content[3:-3].strip()
                plan_data = json.loads(plan_content)

                created_tasks = []
                if isinstance(plan_data, list):
                    for i, task_data in enumerate(plan_data):
                        if isinstance(task_data, dict):
                            task_id = str(uuid.uuid4())
                            task = TaskState(
                                task_id=task_id, status="pending",
                                task_objective=task_data.get("task_objective", f"Planned Objective {i+1}"),
                                description=task_data.get("description", f"Planned Task {i+1}"),
                                selected_agent=task_data.get("selected_agent"),
                                task_inputs=task_data.get("inputs", {}),
                                outputs={}, output_files=[], evaluation={}, error_log=None, feedback_log=None,
                                requires_evaluation=task_data.get("requires_evaluation", False), retry_count=0
                            )
                            if not task["selected_agent"]:
                                print(f"PM Warning (Create): Task {i} lacks 'selected_agent'. Skipping.")
                                continue
                            created_tasks.append(task)
                        else: print(f"PM Warning (Create): Invalid item in plan at index {i}.")
                else:
                    print(f"PM Error (Create): LLM did not return a list. Response: {plan_content[:200]}...")

                if created_tasks:
                    print(f"PM: Initial workflow created with {len(created_tasks)} tasks.")
                    output_state["tasks"] = created_tasks
                    output_state["current_task_index"] = 0
                    output_state["current_task"] = created_tasks[0].copy()
                    tasks = created_tasks # Update local variable
                    current_idx = 0     # Update local variable
                else:
                    print("PM Error (Create): Failed to create valid tasks from initial plan.")
                    output_state["tasks"] = []
                    output_state["current_task_index"] = 0
                    output_state["current_task"] = None
                    return output_state # Let router handle empty tasks
            except Exception as e:
                print(f"PM Error during initial workflow creation: {e}")
                traceback.print_exc()
                output_state["tasks"] = []
                output_state["current_task_index"] = 0
                output_state["current_task"] = None
                return output_state

        # --- Loop to Check Status, Save to LTM, and Handle Failure ---
        print(f"\nPM: --- Entering status check loop ---") # Log entry clearly
        tasks = output_state.get("tasks", []) # Get final task list for the loop
        current_idx = output_state.get("current_task_index", 0) # Get final index for the loop
        print(f"PM: Starting loop with current_idx = {current_idx}, task count = {len(tasks)}")

        while 0 <= current_idx < len(tasks):
            task_to_check = tasks[current_idx]
            status = task_to_check.get("status")
            task_id = task_to_check.get("task_id", "N/A")
            print(f"PM Loop: Checking Task {task_id} (Idx: {current_idx}), Status: '{status}'")

            if status in ["failed", "max_retries_reached"]:
                print(f"PM Loop: Handling '{status}' for task {task_id}")
                await self._save_task_to_ltm(task_to_check) # Save failed task

                # --- MODIFIED RETRY HANDLING ---
                max_retries = _full_static_config.agents.get("process_management", {}).parameters.get("max_retries", 3)
                current_retry_count = task_to_check.get("retry_count", 0)
                fail_action = "SKIP" # Default action if errors occur or retries exceeded

                if current_retry_count >= max_retries:
                    print(f"PM (Fail): Max retries ({max_retries}) already reached for task {task_id}. Forcing SKIP.")
                    fail_action = "SKIP"
                    # Ensure status reflects max retries if it wasn't already
                    if status != "max_retries_reached":
                        tasks[current_idx]["status"] = "max_retries_reached"
                        tasks[current_idx]["feedback_log"] = (tasks[current_idx].get("feedback_log", "") + "\n[PM Note: Marked as max_retries_reached during check.]").strip()
                else:
                    # Retries are available, increment count BEFORE calling LLM
                    current_retry_count += 1
                    tasks[current_idx]["retry_count"] = current_retry_count
                    print(f"PM (Fail): Incremented retry count for task {task_id} to {current_retry_count}/{max_retries}.")

                    # Determine failure type (existing logic)
                    failure_type = "unknown"
                    # ... (existing failure_type determination logic) ...
                    eval_assessment = task_to_check.get("evaluation", {}).get("assessment", "").lower()
                    if eval_assessment == "fail": # Check if eval failed
                        failure_type = "evaluation"
                    elif task_to_check.get("error_log"): # Check for execution error
                        failure_type = "execution"
                    else: # Default if no specific indicators
                        failure_type = "execution" # Or maybe 'unknown'? Let's stick with execution for now.


                    print(f"PM Failure Analysis: Determined Failure Type = '{failure_type}' for Task {task_id}")

                    # Calculate new is_max_retries status AFTER incrementing
                    is_max_retries = current_retry_count >= max_retries

                    failure_context = f"Task failed at index {current_idx}. Status: {status}. Failure Type: {failure_type}." # 加入 type

                    fail_prompt_input = {
                        "failure_context": failure_context,
                        "is_max_retries": is_max_retries, # Pass the updated status
                        "max_retries": max_retries, # Pass the limit itself
                        "selected_agent_name": task_to_check.get("selected_agent", "N/A"),
                        "task_description": task_to_check.get("description", "N/A"),
                        "task_objective": task_to_check.get("task_objective", "N/A"),
                        "inputs_json": json.dumps(task_to_check.get("task_inputs", {}), ensure_ascii=False),
                        "execution_error_log": task_to_check.get("error_log", "None"),
                        "feedback_log": task_to_check.get("feedback_log", "None (Includes Eval Results if any)"),
                        "llm_output_language": self.llm_output_language,
                        # "alternative_agent" is removed based on previous discussion
                        "original_requires_evaluation": task_to_check.get("requires_evaluation", False)
                    }

                    print("--- PM Failure Analysis: Inputs to LLM ---")
                    # ... (log inputs) ...
                    print(f"  Failure Type: {failure_type}")
                    print(f"  Is Max Retries (Passed to LLM): {is_max_retries}") # Log what LLM sees
                    print(f"  Current Retry Count (Passed to LLM): {current_retry_count}") # Log count LLM sees
                    # print(f"  Feedback Log (Input): {fail_prompt_input['feedback_log'][:500]}...")
                    print(f"  Error Log (Input): {fail_prompt_input['execution_error_log']}")
                    print("-----------------------------------------")

                    try:
                        if not self.prompts.get("failure_analysis"): raise ValueError("failure_analysis prompt missing")
                        if not self.llm: raise ValueError("LLM for PM not initialized.")
                        fail_chain = self.llm | StrOutputParser()
                        # Correctly format the prompt with the dictionary
                        fail_content = await fail_chain.ainvoke(self.prompts["failure_analysis"].format(**fail_prompt_input))


                        # Parse LLM response (existing logic)
                        if fail_content.startswith("```json"): fail_content = fail_content[7:-3].strip()
                        elif fail_content.startswith("```"): fail_content = fail_content[3:-3].strip()
                        fail_result = {}
                        try: fail_result = json.loads(fail_content)
                        except json.JSONDecodeError as json_err:
                            print(f"PM Failure Analysis Error: Failed to parse LLM JSON: {json_err}")
                            fail_result = {"action": "SKIP"} # Default to SKIP on parse error

                        # Get action, default to SKIP if missing or invalid
                        fail_action = fail_result.get("action", "SKIP")
                        if fail_action not in ["FALLBACK_GENERAL", "MODIFY", "SKIP"]:
                            print(f"PM Warning (Fail): LLM returned invalid action '{fail_action}'. Defaulting to SKIP.")
                            fail_action = "SKIP"

                        # Log parsed action
                        print(f"--- PM Failure Analysis: Parsed Result ---")
                        print(f"  Parsed Action: {fail_action}")
                        print(f"  Parsed Result Dict: {fail_result}")
                        print("-----------------------------------------")

                        # --- Overwrite LLM action if max retries were reached by the increment ---
                        if is_max_retries and fail_action != "SKIP":
                            print(f"PM (Fail): Max retries reached after increment ({current_retry_count}/{max_retries}). Overriding LLM action '{fail_action}' to SKIP.")
                            fail_action = "SKIP"
                            # Update status if needed
                            if status != "max_retries_reached":
                                tasks[current_idx]["status"] = "max_retries_reached"
                                tasks[current_idx]["feedback_log"] = (tasks[current_idx].get("feedback_log", "") + "\n[PM Note: Marked as max_retries_reached after increment.]").strip()


                    except Exception as e:
                        print(f"PM Error during failure analysis LLM call: {e}")
                        traceback.print_exc()
                        print("PM (Fail): Defaulting to SKIP action due to error during LLM call.")
                        fail_action = "SKIP"
                        # Ensure status reflects failure if possible
                        if status != "max_retries_reached":
                            tasks[current_idx]["status"] = "failed" # Keep as failed if not max retries
                            tasks[current_idx]["error_log"] = (tasks[current_idx].get("error_log", "") + f"\n[PM Error: Failure analysis LLM call failed: {e}]").strip()

                # --- END MODIFIED RETRY HANDLING ---

                # --- Handle Actions (using the potentially modified fail_action) ---
                if fail_action == "FALLBACK_GENERAL":
                    # --- MODIFIED LOG MESSAGE ---
                    print(f"PM (Fail): Processing FALLBACK_GENERAL (ALL new tasks inherit retry count {current_retry_count}).")
                    # --- END MODIFIED LOG ---
                    new_tasks_data = fail_result.get("new_tasks_list")
                    new_task_data = fail_result.get("new_task")
                    created_fallback_tasks = []
                    tasks_source = new_tasks_data if isinstance(new_tasks_data, list) else ([new_task_data] if isinstance(new_task_data, dict) else [])
                    print(f"PM (Fail): FALLBACK_GENERAL expecting {len(tasks_source)} task(s).")

                    for i, task_data in enumerate(tasks_source):
                        if isinstance(task_data, dict):
                            task_id_fb = str(uuid.uuid4())
                            # --- >>> NEW MODIFICATION HERE <<< ---
                            # ALL fallback tasks inherit the same incremented retry count.
                            inherit_retry = current_retry_count
                            # --- >>> END NEW MODIFICATION <<< ---

                            task_fb = TaskState(
                                task_id=task_id_fb, status="pending",
                                task_objective=task_data.get("task_objective", f"Fallback Objective {i+1}"),
                                description=task_data.get("description", f"Fallback Task {i+1}"),
                                selected_agent=task_data.get("selected_agent", "LLMTaskAgent"),
                                task_inputs=task_data.get("inputs", {}),
                                outputs={}, output_files=[], evaluation={}, error_log=None, feedback_log=None,
                                requires_evaluation=task_data.get("requires_evaluation", False),
                                # --- APPLY THE MODIFICATION ---
                                retry_count=inherit_retry # All tasks get the same inherited count
                                # --- END APPLY ---
                            )
                            if not task_fb["selected_agent"]:
                                print(f"PM Warning (Fallback): Task {i} lacks 'selected_agent'. Skipping.")
                                continue
                            created_fallback_tasks.append(task_fb)
                        else: print(f"PM Warning (Fallback): Invalid item {i} in tasks source.")


                    if created_fallback_tasks:
                        # --- MODIFIED LOG MESSAGE ---
                        print(f"PM (Fail): Replacing failed task at index {current_idx} with {len(created_fallback_tasks)} fallback task(s), ALL inheriting retry count {inherit_retry}.")
                        # --- END MODIFIED LOG ---
                        tasks = tasks[:current_idx] + created_fallback_tasks + tasks[current_idx+1:]
                        output_state["tasks"] = tasks
                        output_state["current_task_index"] = current_idx # Index points to the first new task
                        output_state["current_task"] = created_fallback_tasks[0].copy()
                        print(f"PM (Fail): Fallback applied. Continuing loop to process new task at index {current_idx}.")
                        continue # Continue WHILE loop to process the new task
                    else:
                        print("PM Warning (Fail): FALLBACK_GENERAL action chosen, but no valid tasks created/found. Defaulting to SKIP.")
                        fail_action = "SKIP" # If LLM fails to provide tasks, force SKIP

                if fail_action == "MODIFY":
                    print(f"PM (Fail): Modifying task {task_id} based on action 'MODIFY' and resetting to 'pending' for retry.")
                    # Apply modifications suggested by LLM (optional)
                    new_desc = fail_result.get("modify_description")
                    if new_desc and isinstance(new_desc, str):
                        print(f"  - Applying modification: description='{new_desc[:50]}...'")
                        tasks[current_idx]["description"] = new_desc
                    new_obj = fail_result.get("modify_objective")
                    if new_obj and isinstance(new_obj, str):
                        print(f"  - Applying modification: task_objective='{new_obj[:50]}...'")
                        tasks[current_idx]["task_objective"] = new_obj
                    # Add more fields if needed (e.g., modifying task_inputs requires careful handling)

                    # Reset status for retry
                    tasks[current_idx]["status"] = "pending"
                    # Add note to feedback log
                    tasks[current_idx]["feedback_log"] = (tasks[current_idx].get("feedback_log", "") + "\n[PM Note: Modified task based on failure analysis, resetting to pending for retry.]").strip()
                    # Clear error log for the retry attempt
                    tasks[current_idx]["error_log"] = None # Clear error for retry

                    # Update state and break loop
                    output_state["tasks"] = tasks
                    output_state["current_task_index"] = current_idx
                    output_state["current_task"] = tasks[current_idx].copy()
                    print(f"PM (Fail): Task {task_id} modified and set to 'pending'. Breaking loop.")
                    break # Exit WHILE loop, let router handle 'pending'

                # --- Handle SKIP Action (handles forced skip and LLM skip) ---
                if fail_action == "SKIP":
                    print(f"PM (Fail): Skipping task {task_id} at index {current_idx} based on action 'SKIP'.")
                    # Ensure status is max_retries_reached if it was forced skip due to retries
                    if tasks[current_idx].get("retry_count", 0) >= max_retries:
                         tasks[current_idx]["status"] = "max_retries_reached"
                         # --- FIX: Save skipped task to LTM ---
                         print(f"PM Loop: Saving skipped/max_retries task {task_id} to LTM before skipping.")
                         await self._save_task_to_ltm(tasks[current_idx])
                         # --- END FIX ---

                    current_idx += 1
                    output_state["current_task_index"] = current_idx # Update index before continuing
                    # Ensure current_task reflects the next task or None
                    output_state["current_task"] = tasks[current_idx].copy() if 0 <= current_idx < len(tasks) else None
                    print(f"PM (Fail): Updated current_idx to {current_idx}. Continuing loop.")
                    continue # Continue WHILE loop to check the NEXT task

            # --- End of failed/max_retries_reached block ---

            elif status == "completed":
                # ... (原本的 completed handling 邏輯) ...
                # --- 確保增加 current_idx 並 continue ---
                 print(f"PM Loop: Handling 'completed' for task {task_id}. Saving to LTM.")
                 await self._save_task_to_ltm(task_to_check)
                 current_idx += 1 # Increment index AFTER processing the completed task
                 # --- Update state *after* incrementing ---
                 output_state["current_task_index"] = current_idx # Update index in state
                 # Update current_task based on new index before continuing loop check
                 output_state["current_task"] = tasks[current_idx].copy() if 0 <= current_idx < len(tasks) else None
                 print(f"PM Loop: Incremented index to {current_idx} after completing task {task_id}. Continuing loop check.")
                 continue # Go back to the start of the while loop check


            elif status in ["pending", "in_progress", None]:
                # 發現需要執行的任務，設置狀態並跳出循環
                print(f"PM Loop: Task {task_id} is '{status}'. Ready for routing. Breaking loop.")
                output_state["current_task_index"] = current_idx
                output_state["current_task"] = task_to_check.copy()
                break # 跳出 WHILE 循環

            else: # Unexpected status
                # ... (原本的 unexpected status handling 邏輯) ...
                # --- 確保增加 current_idx 並 continue ---
                print(f"PM Loop Warning: Unhandled task status '{status}' for Task {task_id}. Skipping.")
                await self._save_task_to_ltm(task_to_check)
                current_idx += 1
                output_state["current_task_index"] = current_idx
                output_state["current_task"] = tasks[current_idx].copy() if 0 <= current_idx < len(tasks) else None
                print(f"PM Loop (Warning): Updated current_idx to {current_idx}. Continuing loop.")
                continue


        # --- After the loop ---
        # Logging remains the same
        print(f"PM: --- Debug log: Exited status check loop ---")
        final_idx_after_loop = output_state.get("current_task_index", current_idx) # Get potentially updated index
        print(f"PM: Final index after loop processing: {final_idx_after_loop}")
        tasks_final_count = len(output_state.get("tasks", [])) # Get final task count
        if final_idx_after_loop >= tasks_final_count:
            print("PM: Reached end of task list after loop.")
            # State for current_task/index already set inside loop or at end
            output_state["current_task"] = None # Ensure current task is None if index is out of bounds
            output_state["current_task_index"] = final_idx_after_loop # Keep the final index
        else:
             # Task found, index/task already set inside loop's break condition
             # Ensure state reflects this correctly
             output_state["current_task_index"] = final_idx_after_loop
             if "current_task" not in output_state or output_state["current_task"] is None or output_state["current_task"].get("task_id") != tasks[final_idx_after_loop].get("task_id"):
                  output_state["current_task"] = tasks[final_idx_after_loop].copy()
             print(f"PM: Exiting with index {final_idx_after_loop} pointing to task {output_state['current_task'].get('task_id') if output_state.get('current_task') else 'N/A'}.")

        print(f"=== Process Management Finished. Routing Index: {output_state.get('current_task_index', 'N/A')}, Phase: {output_state.get('current_phase')} ===")
        return output_state # 返回最終狀態

class EvaAgent:
    """
    Parent node responsible for initiating evaluation by invoking the evaluation_subgraph.
    It now acts as a simple trigger based on routing.
    """
    def __init__(self):
        pass

    async def run(self, state: WorkflowState, config: RunnableConfig) -> Dict[str, Any]:
        """
        Checks if evaluation is needed (based on routing) and invokes subgraph.
        Evaluation type logic is handled within the subgraph based on selected_agent.
        """
        print("--- Running EvaAgent Node (Simple Trigger) ---")
        current_idx = state.get("current_task_index") # Get from state directly
        original_tasks_from_state = state.get("tasks", [])
        
        # Create a deep copy of tasks to modify locally if needed, but prefer using subgraph's output
        tasks = [t.copy() for t in original_tasks_from_state]

        if current_idx is None or not (0 <= current_idx < len(tasks)):
            print(f"EvaAgent Error: Invalid task index '{current_idx}'. Tasks length: {len(tasks)}.")
            # Return the original state or an empty dict if critical
            return {"tasks": original_tasks_from_state} # Return original tasks to avoid state corruption

        current_task = tasks[current_idx] # This is a copy

        # Simplified check: If routed here, assume evaluation is needed and pending
        # The status and requires_evaluation should be checked by ProcessManagementAgent before routing here
        print(f"EvaAgent: Triggering evaluation for Task {current_task.get('id', 'N/A')} (Agent: {current_task.get('selected_agent')}, Original Task Status: {current_task.get('status')})")

        # --- Set status to in_progress on our local copy for the subgraph invocation ---
        # The subgraph will operate on a state that reflects this.
        # The *actual* task update will come from the subgraph's output.
        current_task_for_subgraph = current_task.copy()
        current_task_for_subgraph["status"] = "in_progress" # Use string literal

        # Prepare the state to be passed to the subgraph
        # Ensure we're passing a state object that the subgraph expects
        subgraph_input_state_dict = state.copy() # Create a shallow copy of the input state dictionary
        
        # Update the tasks list within this copied state for the subgraph
        temp_tasks_for_subgraph = [t.copy() for t in original_tasks_from_state]
        if 0 <= current_idx < len(temp_tasks_for_subgraph):
            temp_tasks_for_subgraph[current_idx] = current_task_for_subgraph
        subgraph_input_state_dict["tasks"] = temp_tasks_for_subgraph
        
        # --- Invoke Subgraph ---
        try:
            print(f"EvaAgent: Invoking evaluation subgraph with task '{current_task_for_subgraph.get('id')}' set to in_progress...")
            
            # The subgraph `evaluation_teams` is compiled to operate on `WorkflowState` (or its dict form).
            # It should return the full, updated state.
            subgraph_output_state_dict: Dict[str, Any] = await evaluation_teams.ainvoke(
                subgraph_input_state_dict, # Pass the modified state dict
                config=config
            )
            print(f"EvaAgent: Evaluation subgraph finished.")

            # --- Process Subgraph Output ---
            # The subgraph_output_state_dict IS the new state. We need to extract 'tasks' from it.
            final_tasks_from_subgraph = subgraph_output_state_dict.get("tasks")

            if final_tasks_from_subgraph is None:
                print("EvaAgent Error: Subgraph did not return 'tasks' in its output state.")
                # Fallback: Mark the original task as failed
                tasks[current_idx]["status"] = "failed" # Use string literal
                tasks[current_idx]["error_log"] = (tasks[current_idx].get("error_log", "") + "; EvaAgent: Subgraph output missing 'tasks'").strip("; ")
                tasks[current_idx]["feedback_log"] = (tasks[current_idx].get("feedback_log", "") + "; EvaAgent: Subgraph output missing 'tasks'").strip("; ")
                return {"tasks": tasks} # Return the locally modified tasks

            # Log details of the specific task that was evaluated
            if 0 <= current_idx < len(final_tasks_from_subgraph):
                final_task_after_subgraph = final_tasks_from_subgraph[current_idx]
                print(f"Task {final_task_after_subgraph.get('id', 'N/A')} (Index: {current_idx}) after evaluation subgraph:")
                print(f"  - Status: {final_task_after_subgraph.get('status')}")
                print(f"  - Outputs type: {type(final_task_after_subgraph.get('task_outputs'))}")
                if isinstance(final_task_after_subgraph.get('task_outputs'), dict):
                    print(f"  - Outputs assessment: {final_task_after_subgraph['task_outputs'].get('assessment', 'N/A')}")
                    print(f"  - Outputs feedback snippet: {str(final_task_after_subgraph['task_outputs'].get('feedback', 'N/A'))[:100]}...")
                print(f"  - Feedback Log snippet: {str(final_task_after_subgraph.get('feedback_log', 'N/A'))[:100]}...")
                print(f"  - Error Log: {final_task_after_subgraph.get('error_log', 'N/A')}")
            else:
                print(f"EvaAgent Warning: current_idx '{current_idx}' out of bounds for final_tasks_from_subgraph (len: {len(final_tasks_from_subgraph)}).")


            # The main output of EvaAgent should be the new state of 'tasks'
            return {"tasks": final_tasks_from_subgraph}

        except Exception as e:
            print(f"EvaAgent: Critical error invoking or processing evaluation subgraph: {e}")
            import traceback
            traceback.print_exc()
            
            # Fallback: Mark the original task (from the input state to EvaAgent) as failed
            # This 'tasks' is the copy made at the beginning of this method.
            if 0 <= current_idx < len(tasks):
                task_to_fail = tasks[current_idx]
                task_to_fail["status"] = "failed" # Use string literal
                error_message = f"Evaluation Subgraph Invocation Error: {type(e).__name__} - {e}"
                
                existing_feedback_log = task_to_fail.get("feedback_log", "")
                task_to_fail["feedback_log"] = (existing_feedback_log + "; " + error_message).strip("; ")
                
                existing_error_log = task_to_fail.get("error_log", "")
                task_to_fail["error_log"] = (existing_error_log + "; " + error_message).strip("; ")

                if "evaluation" not in task_to_fail or not isinstance(task_to_fail["evaluation"], dict):
                    task_to_fail["evaluation"] = {}
                task_to_fail["evaluation"]["assessment"] = "Fail" # Consistent with image
                task_to_fail["evaluation"]["subgraph_error"] = error_message
                
                print(f"--- EvaAgent Exception Handling ---")
                print(f"  Error: {error_message}")
                print(f"  Marked Task ID: {task_to_fail.get('id', 'N/A')} as FAILED.")
                print(f"  Task Status: {task_to_fail.get('status')}")
                print(f"  Task Feedback Log: {task_to_fail.get('feedback_log')}")
                print(f"  Task Error Log: {task_to_fail.get('error_log')}")
                print("---------------------------------")
            else:
                print("EvaAgent Error: current_idx invalid during exception handling. Cannot mark specific task as failed.")

            return {"tasks": tasks} # Return the locally modified tasks with the error


class QA_Agent:
    def __init__(self, long_term_memory_vectorstore=vectorstore):
        self.vectorstore = long_term_memory_vectorstore
        # --- 從配置加載模板字串 ---
        qa_prompt_config = _qa_prompts.get("qa_prompt") # 獲取 PromptConfig 對象
        self.qa_prompt_template_str = qa_prompt_config.template if qa_prompt_config else "" # 獲取模板字串
        self.qa_prompt_input_variables = qa_prompt_config.input_variables if qa_prompt_config else [] # 獲取變數列表

        if not self.qa_prompt_template_str:
             print("QA_Agent WARNING: qa_prompt template string not found in config. Using a basic default.")
             # 使用一個非常基礎的預設模板，以防萬一
             self.qa_prompt_template_str = """對話記錄:\n{chat_history}\n\n人類: {last_user_query}\nAI:"""
             self.qa_prompt_input_variables = ["chat_history", "last_user_query"] # 預設模板的變數

        # --- 驗證必要的變數是否存在 ---
        # 根據 configuration.py 的預期變數進行檢查
        expected_vars = {"task_summary", "retrieved_ltm_context", "chat_history", "last_user_query", "llm_output_language"}
        if not expected_vars.issubset(self.qa_prompt_input_variables):
            print(f"QA_Agent WARNING: Loaded qa_prompt template might be missing expected variables.")
            print(f"  Expected: {expected_vars}")
            print(f"  Found in config: {self.qa_prompt_input_variables}")
            # Consider raising an error or falling back more gracefully

        # --- STM setup ---
        self.memory_window_size = _full_static_config.agents.get("qa_agent", {}).parameters.get("memory_window_size", 10)
        # Note: ConversationBufferWindowMemory formats history into a single string key, matching 'chat_history'
        self.stm = ConversationBufferWindowMemory(
            k=self.memory_window_size, return_messages=False, # Set return_messages=False to get formatted string
            memory_key="chat_history", input_key="human_input" # Use different keys for clarity
        )
        # --- 創建 ChatPromptTemplate 實例 ---
        # 使用從配置中讀取的模板字串
        self.prompt_template = ChatPromptTemplate.from_template(self.qa_prompt_template_str)
        print(f"QA_Agent initialized using prompt template from configuration: 'qa_prompt'")


    async def run(self, state: WorkflowState, config: RunnableConfig) -> Dict[str, Any]:
        node_name = "QA Agent Node (chat_bot)"
        print(f"--- 運行節點: {node_name} ---")

        # 獲取配置和上下文
        runtime_config = config["configurable"]
        qa_llm_config_params = {
            "model_name": runtime_config.get("qa_model_name"),
            "temperature": runtime_config.get("qa_temperature"),
            "max_tokens": runtime_config.get("qa_max_tokens"),
        }
        qa_llm_config_params = {k: v for k, v in qa_llm_config_params.items() if v is not None}
        llm = initialize_llm(qa_llm_config_params, agent_name_for_default_lookup="qa_agent")
        
        retriever_k = runtime_config.get("retriever_k", 5)
        retriever = self.vectorstore.as_retriever(search_kwargs=dict(k=retriever_k))
        ltm = VectorStoreRetrieverMemory(retriever=retriever, memory_key=ltm_memory_key, input_key=ltm_input_key)
        llm_output_language = runtime_config.get("global_llm_output_language", LLM_OUTPUT_LANGUAGE_DEFAULT)

        last_user_query_content = state.get("interrupt_input")
        
        # 初始化返回狀態
        # interrupt_input 預設會被 QA Agent 消耗並清除。
        # 只有在特定情況下 (如 RESUME_TASK 且帶有附加指令) 才可能被重新填充。
        return_state_update = {
            "interrupt_input": None 
        }

        try:
            if last_user_query_content:
                print(f"{node_name}: 直接從 interrupt_input 處理使用者查詢: '{last_user_query_content[:100]}...'")
                query_for_ltm = last_user_query_content

                retrieved_ltm_context = "LTM: N/A"
                if query_for_ltm:
                    try:
                        ltm_loaded_vars = await ltm.aload_memory_variables({ltm_input_key: query_for_ltm})
                        context_from_ltm = ltm_loaded_vars.get(ltm.memory_key)
                        if context_from_ltm:
                            retrieved_ltm_context = context_from_ltm
                    except Exception as e:
                        print(f"{node_name}: LTM錯誤: {e}")

                task_summary = "\n".join([f"- Task {i+1}: {t['description']} ({t['status']})" for i, t in enumerate(state.get("tasks", []))]) or "沒有執行任何任務。"
                current_stm_vars = self.stm.load_memory_variables({})
                formatted_chat_history = current_stm_vars.get(self.stm.memory_key, "沒有STM歷史。")

                chain_input = {
                    "last_user_query": last_user_query_content,
                    "retrieved_ltm_context": retrieved_ltm_context,
                    "window_size": self.memory_window_size,
                    "chat_history": formatted_chat_history,
                    "llm_output_language": llm_output_language,
                    "task_summary": task_summary,
                }

                qa_chain = self.prompt_template | llm
                response_message = await qa_chain.ainvoke(chain_input)
                response_content = response_message.content.strip()
                print(f"{node_name} 原始回應: {response_content[:150]}...")

                next_phase = "qa" # Default to staying in QA
                terminate_indicators = ["TERMINATE", "結束對話", "再見", "謝謝再見", "不需要了"]
                is_terminate = any(ind in response_content for ind in terminate_indicators)
                resume_indicators = ["RESUME_TASK", "繼續任務", "返回任務", "回到任務", "繼續工作流程"]
                is_resume = any(ind in response_content for ind in resume_indicators)
                new_task_match = None
                if response_content.startswith("NEW_TASK:"):
                    new_task_match = response_content[len("NEW_TASK:") :].strip()

                if is_terminate:
                    response_content = "好的，對話結束。"
                    next_phase = "finished"
                    print(f"{node_name}: 檢測到終止對話意圖")
                elif is_resume:
                    response_content = "好的，正在返回任務執行流程。"
                    next_phase = "task_execution"
                    # --- MODIFICATION FOR RESUME_TASK ---
                    # 將使用者在要求 RESUME_TASK 時的原始查詢 (last_user_query_content) 
                    # 重新放入 interrupt_input，以便 PM 可以將其作為插隊請求來處理。
                    print(f"{node_name}: Detected RESUME_TASK. Passing original query '{last_user_query_content}' as interrupt_input to PM for potential interrupt processing.")
                    return_state_update["interrupt_input"] = last_user_query_content
                    # state["user_input"] (for PM initial planning) should not be set here,
                    # as this is an interrupt to an existing flow, not a brand new workflow request.
                    # --- END MODIFICATION ---
                elif new_task_match:
                    response_content = f"收到新任務：'{new_task_match}'。正在返回任務規劃..."
                    next_phase = "task_execution"
                    print(f"{node_name}: Detected NEW_TASK. Clearing existing tasks and setting new user_input for PM.")
                    return_state_update["user_input"] = new_task_match
                    return_state_update["tasks"] = []
                    return_state_update["current_task_index"] = 0
                    return_state_update["current_task"] = None
                else:
                    # Standard QA response, stay in QA phase
                    print(f"{node_name}: 普通回答，維持QA階段")
                    next_phase = "qa" # Explicitly ensure phase is qa for normal chat

                self.stm.save_context({"human_input": last_user_query_content}, {"output": response_content})
                
                return_state_update["current_phase"] = next_phase
                current_qa_context = [AIMessage(content=response_content)]
                return_state_update["qa_context"] = current_qa_context
                
                print(f"{node_name}: 返回狀態: phase='{next_phase}', 消息='{response_content[:50]}...'")
                return return_state_update
            
            else: # last_user_query_content (來自 interrupt_input) 為空
                print(f"{node_name}: 本輪無新的 interrupt_input。提供通用提示。")
                
                prompt_text = ""
                # 檢查 STM 是否為空，以判斷是初次進入還是AI等待輸入
                if not self.stm.chat_memory.messages: # STM 為空，代表首次進入 QA
                    prompt_text = "您好，請問有什麼可以協助您的嗎？ (您可以輸入 '結束對話' 或 '繼續任務')"
                else: # STM 不為空，表示之前有過對話，AI 現在等待使用者輸入
                    prompt_text = "還有其他問題嗎？或者請指示 '繼續任務' 或 '結束對話'。"
                
                print(f"{node_name}: 發送提示: {prompt_text}")
                return_state_update["current_phase"] = "qa"
                return_state_update["qa_context"] = [AIMessage(content=prompt_text)]
                # 此引導性提示不儲存到 STM，因為它不是對使用者具體查詢的回應
                return return_state_update

        except Exception as e:
            print(f"{node_name} 錯誤: {e}")
            traceback.print_exc()
            error_message = f"處理您的問題時發生錯誤: {e}"
            return_state_update["current_phase"] = "qa"
            return_state_update["qa_context"] = [AIMessage(content=error_message)]
            return return_state_update


# --- Keep loading agent descriptions as they are still needed ---
agent_descriptions = _full_static_config.agents.get("assign_agent", {}).parameters.get("specialized_agents_description", {})
if not agent_descriptions:
     print("WARNING: specialized_agents_description not found in config for AssignAgent!")

# =============================================================================
# 7. Agent 實例化
# =============================================================================
process_management = ProcessManagement(long_term_memory_vectorstore=vectorstore)
eva_agent = EvaAgent() # Instantiate the refactored EvaAgent
qa_agent = QA_Agent(long_term_memory_vectorstore=vectorstore)


def parse_llm_json_output(llm_output: str) -> dict:
    """
    穩健地解析來自LLM的JSON輸出，處理潛在的Markdown程式碼區塊標記。

    Args:
        llm_output: LLM返回的原始字串。

    Returns:
        一個包含解析後JSON數據的字典，或者一個包含錯誤訊息的字典。
    """
    if not isinstance(llm_output, str):
        return {"error": "LLM output is not a string.", "raw_output": str(llm_output)}

    text_to_parse = llm_output.strip()
    json_str = None

    # 優先嘗試匹配 ```json ... ```
    match_json_block = re.search(r"```json\s*([\s\S]*?)\s*```", text_to_parse, re.DOTALL)
    if match_json_block:
        json_str = match_json_block.group(1).strip()
    else:
        # 其次嘗試匹配 ``` ... ``` (不帶json標記)
        match_any_block = re.search(r"```\s*([\s\S]*?)\s*```", text_to_parse, re.DOTALL)
        if match_any_block:
            json_str = match_any_block.group(1).strip()
        else:
            # 如果沒有找到程式碼區塊標記，則假定整個文本是JSON
            json_str = text_to_parse

    # 檢查提取的字串是否看起來像一個JSON對象或數組
    # 這是一個啟發式檢查，因為有時json_str可能是空字串或不包含有效的JSON內容
    if not json_str or not ( (json_str.startswith('{') and json_str.endswith('}')) or \
                             (json_str.startswith('[') and json_str.endswith(']')) ):
        # 如果剝離標記後不像JSON，或者json_str為空，嘗試解析原始輸入 `text_to_parse`
        # 因為可能LLM返回了不帶標記的純JSON，或者是一個錯誤訊息
        try:
            parsed_json = json.loads(text_to_parse)
            if isinstance(parsed_json, dict): # 我們期望Sankey結構是字典
                 return parsed_json
            # 如果解析成功但不是字典 (例如，LLM返回了一個純字串錯誤訊息且恰好是有效的JSON字串)
            # 則認為這不是我們期望的Sankey結構
        except json.JSONDecodeError:
            # 如果原始文本也無法解析，並且json_str是無效的，則返回錯誤
            if not json_str or not ( (json_str.startswith('{') and json_str.endswith('}')) or \
                                     (json_str.startswith('[') and json_str.endswith(']')) ):
                return {"error": "Failed to extract a valid JSON structure from LLM output.", "processed_string": json_str, "raw_output": llm_output}
            # 否則，繼續嘗試解析 json_str
    
    try:
        parsed_json = json.loads(json_str)
        # 我們期望Sankey的結構是一個字典
        if isinstance(parsed_json, dict):
            return parsed_json
        else:
            return {"error": f"Parsed JSON is not a dictionary. Type: {type(parsed_json)}", "processed_string": json_str, "raw_output": llm_output}
    except json.JSONDecodeError as e:
        # 如果標準解析失敗，嘗試更積極地提取第一個 '{' 和最後一個 '}' 之間的內容
        # 這有助於處理JSON前後可能存在的少量雜訊文本（即使在```之外）
        try:
            start_brace = json_str.find('{')
            end_brace = json_str.rfind('}')
            if start_brace != -1 and end_brace != -1 and end_brace > start_brace:
                potential_json_str = json_str[start_brace : end_brace+1]
                parsed_json_aggressive = json.loads(potential_json_str)
                if isinstance(parsed_json_aggressive, dict):
                    return parsed_json_aggressive
                else: # 積極解析成功但不是字典
                    return {"error": f"Aggressively parsed JSON is not a dictionary. Type: {type(parsed_json_aggressive)}", "processed_string": potential_json_str, "raw_output": llm_output}
            else: # 未找到大括號或順序無效
                return {"error": f"JSONDecodeError: {e}. Could not find valid JSON object structure.", "processed_string": json_str, "raw_output": llm_output}
        except json.JSONDecodeError as e2: # 積極提取後仍然解析失敗
            return {"error": f"JSONDecodeError even after aggressive extraction: {e2}", "processed_string": json_str, "raw_output": llm_output}
    except Exception as e_general: # 捕獲其他意外錯誤
        return {"error": f"An unexpected error occurred during JSON parsing: {str(e_general)}", "processed_string": json_str, "raw_output": llm_output}


# --- <<< NEW: SankeyStructureAgent Class Definition >>> ---
class SankeyStructureAgent:
    async def run(self, state: WorkflowState, config: RunnableConfig) -> WorkflowState:
        print("--- Running SankeyStructureAgent ---")
        output_state = state.copy()
        
        try:
            runtime_cfg_dict = config.get("configurable", {})
            user_input_for_sankey = state.get("user_input", "N/A")
            tasks_for_sankey_original = state.get("tasks", [])

            # --- NEW: Clean tasks before sending to LLM ---
            cleaned_tasks_for_llm = []
            for task_item_dict in tasks_for_sankey_original:
                if not isinstance(task_item_dict, dict):
                    # Should not happen with TaskState, but good to be safe
                    cleaned_tasks_for_llm.append(str(task_item_dict)) 
                    continue

                current_cleaned_task = {}
                for t_key, t_val in task_item_dict.items():
                    if t_key == "output_files" and isinstance(t_val, list):
                        current_cleaned_task[t_key] = []
                        for f_item in t_val:
                            if isinstance(f_item, dict):
                                f_copy = f_item.copy()
                                f_copy.pop("base64_data", None)  # Remove base64 data
                                # Optionally, remove other large or less relevant fields from file info for LLM
                                # f_copy.pop("path", None) # Example: if path is too long or not needed
                                current_cleaned_task[t_key].append(f_copy)
                            else:
                                current_cleaned_task[t_key].append(f_item)
                    elif t_key == "outputs" and isinstance(t_val, dict):
                        # Optionally, clean large text fields within outputs if they exist and are problematic
                        # For now, focusing on base64 as per user feedback
                        # Example: Prune very long string values in outputs
                        cleaned_outputs = {}
                        for out_key, out_val in t_val.items():
                            if isinstance(out_val, str) and len(out_val) > 1000: # Arbitrary limit
                                cleaned_outputs[out_key] = out_val[:1000] + "... (truncated)"
                            else:
                                cleaned_outputs[out_key] = out_val
                        current_cleaned_task[t_key] = cleaned_outputs
                    elif t_key == "task_inputs" and isinstance(t_val, dict) : # Also check inputs
                        cleaned_inputs = {}
                        for in_key, in_val in t_val.items():
                            if isinstance(in_val, str) and len(in_val) > 1000:
                                cleaned_inputs[in_key] = in_val[:1000] + "... (truncated)"
                            else:
                                cleaned_inputs[in_key] = in_val
                        current_cleaned_task[t_key] = cleaned_inputs
                    # Keep other serializable fields as is
                    elif isinstance(t_val, (dict, list, str, int, float, bool, type(None))):
                        current_cleaned_task[t_key] = t_val
                    else:
                        # For non-serializable, represent as string (should ideally not happen with TypedDicts)
                        current_cleaned_task[t_key] = f"<Non-serializable type {type(t_val).__name__}>"
                cleaned_tasks_for_llm.append(current_cleaned_task)
            
            full_tasks_json_str = json.dumps(cleaned_tasks_for_llm, ensure_ascii=False, default=str)
            # --- END NEW ---

            # --- Get LLM Config for Sankey Structure Agent ---
            agent_name_key = "sankey_structure_agent"
            sankey_agent_llm_config_params = {
                "model_name": runtime_cfg_dict.get(f"sankey_structure_model_name"),
                "temperature": runtime_cfg_dict.get(f"sankey_structure_temperature"),
                "max_tokens": runtime_cfg_dict.get(f"sankey_structure_max_tokens"),
            }
            # Filter out None values if not present in runtime_config
            sankey_agent_llm_config_params = {k: v for k, v in sankey_agent_llm_config_params.items() if v is not None}
            
            # Initialize LLM
            sankey_llm = initialize_llm(sankey_agent_llm_config_params, agent_name_for_default_lookup=agent_name_key)

            # --- Get Prompt Template ---
            # Try runtime config first, then base default
            sankey_prompt_template_str = runtime_cfg_dict.get(
                "sankey_structure_agent_prompt", 
                _full_static_config.agents.get(agent_name_key, {}).prompts.get("generate_sankey_structure", {}).template # MODIFIED HERE
            )
            if not sankey_prompt_template_str:
                print("SankeyStructureAgent ERROR: Prompt template for generate_sankey_structure not found.")
                output_state["sankey_structure_data"] = {"error": "Prompt template missing"}
                return output_state

            prompt = PromptTemplate.from_template(sankey_prompt_template_str)
            
            # --- Prepare Chain ---
            # MODIFIED: Use StrOutputParser to get raw string, then our robust parser
            string_parser = StrOutputParser()
            chain = prompt | sankey_llm | string_parser

            # --- Invoke LLM ---
            print(f"SankeyStructureAgent: Invoking LLM with {len(cleaned_tasks_for_llm)} cleaned tasks.") # Modified log
            
            llm_output_language = runtime_cfg_dict.get("global_llm_output_language", LLM_OUTPUT_LANGUAGE_DEFAULT)

            # MODIFIED: Get raw string output
            llm_raw_output_str = await chain.ainvoke({
                "user_input": user_input_for_sankey,
                "full_tasks_json": full_tasks_json_str, 
                "llm_output_language": llm_output_language 
            })
            print("SankeyStructureAgent: LLM call successful, received raw string output.")
            # print(f"  Raw LLM Output String for parsing: {llm_raw_output_str[:500]}...") # For debugging

            # MODIFIED: Parse the raw string output robustly
            sankey_structure_json = parse_llm_json_output(llm_raw_output_str)


            # Validate the structure (basic check)
            if isinstance(sankey_structure_json, dict) and \
               "nodes" in sankey_structure_json and isinstance(sankey_structure_json["nodes"], list) and \
               "links" in sankey_structure_json and isinstance(sankey_structure_json["links"], list):
                output_state["sankey_structure_data"] = sankey_structure_json
                print(f"SankeyStructureAgent: Successfully processed and stored Sankey structure data ({len(sankey_structure_json['nodes'])} nodes, {len(sankey_structure_json['links'])} links).")
            else:
                # MODIFIED: Improved error handling for parse_llm_json_output results
                error_message_detail = "LLM output was not in the expected format (missing nodes/links lists or parsing error)."
                
                if isinstance(sankey_structure_json, dict) and "error" in sankey_structure_json:
                    # Error came directly from parse_llm_json_output
                    # The 'error' value from parse_llm_json_output might already be quite descriptive.
                    # We can log more details if available in the error dict from the parser.
                    print(f"SankeyStructureAgent ERROR: Failed to parse LLM output. Parser error: '{sankey_structure_json.get('error', 'Unknown parsing error')}'")
                    if "processed_string" in sankey_structure_json:
                         print(f"  Attempted to parse: {sankey_structure_json['processed_string'][:500]}...")
                    if "raw_output" in sankey_structure_json:
                         print(f"  Original raw LLM output: {sankey_structure_json['raw_output'][:500]}...")
                    output_state["sankey_structure_data"] = sankey_structure_json # Pass the detailed error dict
                else:
                    # Parsed successfully into a dict by parse_llm_json_output, but not the right Sankey structure
                    print(f"SankeyStructureAgent ERROR: {error_message_detail} Parsed output: {sankey_structure_json}")
                    output_state["sankey_structure_data"] = {"error": error_message_detail, "llm_output_after_parse": sankey_structure_json, "raw_llm_str": llm_raw_output_str}
        
        except Exception as e:
            print(f"SankeyStructureAgent ERROR: An exception occurred in the agent run: {e}")
            traceback.print_exc()
            # Ensure a consistent error structure for generate_sankey_diagram_plotly
            output_state["sankey_structure_data"] = {"error": f"Agent execution exception: {str(e)}"}
            
        return output_state

sankey_structure_agent = SankeyStructureAgent() # Instantiate the new agent
# --- <<< END NEW Agent Definition and Instantiation >>> ---

# =============================================================================
# 8. 輔助/節點函數定義
# =============================================================================
# Helper function for Sankey diagram node labels
# def task_agent_short_name(agent_name: Optional[str]) -> str:
#     if not agent_name:
#         return "Unknown" # Shorter
#     if "LLMTaskAgent" in agent_name: return "ConceptGen" # More specific to its typical first role
#     if "ImageGenerationAgent" in agent_name: return "ImgGen"
#     if "SpecialEvaAgent" in agent_name: return "BranchEval" # More descriptive of its role
#     if "FinalEvaAgent" in agent_name: return "FinalSelect" # More descriptive
#     if "RhinoMCPCoordinator" in agent_name: return "Rhino"
#     if "PinterestMCPCoordinator" in agent_name: return "Pin"
#     if "OSMMCPCoordinator" in agent_name: return "OSM"
#     if "ArchRAGAgent" in agent_name: return "RAG"
#     if "ImageRecognitionAgent" in agent_name: return "ImgRec"
#     if "VideoRecognitionAgent" in agent_name: return "VidRec"
#     if "ModelRenderAgent" in agent_name: return "Render"
#     if "Generate3DAgent" in agent_name: return "3DGen"
#     if "WebSearchAgent" in agent_name: return "Web"
#     if "EvaAgent" in agent_name: return "StdEval" # Standard Evaluation
    
#     # General fallback
#     name = agent_name.replace("Agent", "").replace("Coordinator", "")
#     if len(name) > 10: # Try to abbreviate further
#         parts = name.split("MCP")
#         if len(parts) > 1 and parts[0]:
#             return parts[0][:7] + "MCP" 
#         return name[:10]
#     return name


def generate_sankey_diagram_plotly(
    sankey_structure_data: Optional[Dict[str, List[Dict]]], 
    user_input_for_title: str, # Changed from full state to just necessary parts
    output_dir: str,
    num_tasks_for_width_heuristic: int = 0 # Added for width heuristic
) -> Optional[str]:
    """
    Generates a Sankey diagram using Plotly based on pre-structured node and link data.
    """
    print("--- Generating Sankey Diagram (v5 - LLM Structured with Morandi) ---")

    if not sankey_structure_data or \
       not isinstance(sankey_structure_data, dict) or \
       "nodes" not in sankey_structure_data or \
       "links" not in sankey_structure_data or \
       not isinstance(sankey_structure_data["nodes"], list) or \
       not isinstance(sankey_structure_data["links"], list):
        print("Sankey Diagram (v5): Invalid or missing sankey_structure_data. Skipping generation.")
        if sankey_structure_data and "error" in sankey_structure_data:
             print(f"  Error from SankeyStructureAgent: {sankey_structure_data['error']}")
        return None

    structured_nodes = sankey_structure_data["nodes"]
    structured_links = sankey_structure_data["links"]

    if not structured_nodes or not structured_links:
        print("Sankey Diagram (v5): No nodes or links provided in the structure data. Skipping generation.")
        return None

    node_indices_map = {}
    plotly_node_labels = []
    plotly_node_colors = []
    
    # --- MODIFIED COLORS (Morandi Palette) ---
    morandi_colors = [
        '#686789', '#B77F70', '#E5E2B9', '#BEB1A8', '#A79A89', '#8A95A9',
        '#ECCED0', '#7D7465', '#E8D3C0', '#7A8A71', '#789798', '#B57C82',
        '#9FABB9', '#B0B1B6', '#99857E', '#88878D', '#91A0A5', '#9AA690'
    ]
    # Ensure enough colors by cycling through morandi_colors if needed
    num_morandi = len(morandi_colors)

    color_map_by_type = {
        "workflow_control": morandi_colors[0 % num_morandi], 
        "task_box": morandi_colors[1 % num_morandi],    
        "concept_item": morandi_colors[2 % num_morandi],
        "image_item": morandi_colors[3 % num_morandi],  
        "rhino_item": morandi_colors[4 % num_morandi],  
        "selected_item": morandi_colors[5 % num_morandi],
        "default_node": morandi_colors[6 % num_morandi],
        "evaluation_summary_item": morandi_colors[16 % num_morandi] # Added for eval summary
    }
    agent_color_map = { # These can override task_box if agent_short_name is provided
        "ConceptGen": morandi_colors[7 % num_morandi],   
        "ImgGen": morandi_colors[8 % num_morandi],      
        "BranchEval": morandi_colors[9 % num_morandi], 
        "FinalSelect": morandi_colors[10 % num_morandi], 
        "Rhino": morandi_colors[11 % num_morandi],  
        "Pin": morandi_colors[12 % num_morandi],
        "OSM": morandi_colors[13 % num_morandi],
        "RAG": morandi_colors[14 % num_morandi],
        "StdEval": morandi_colors[15 % num_morandi],
        # Add more agents if their short names are used by LLM
    }
    default_node_color = color_map_by_type["default_node"]
    # --- END MODIFIED COLORS ---


    for i, node_data in enumerate(structured_nodes):
        node_id = node_data.get("id")
        if not node_id:
            print(f"Sankey Diagram (v5) Warning: Node at index {i} missing 'id'. Skipping.")
            continue
        
        node_label = node_data.get("label", node_id) # Get the label
        
        # --- ADDED: Attempt to sanitize node label encoding for PNG export ---
        if isinstance(node_label, str):
            try:
                # Encode to utf-8 and then decode, ignoring errors.
                # This removes characters that cannot be represented in utf-8.
                sanitized_label = node_label.encode('utf-8', 'ignore').decode('utf-8')
                 # Also replace common problematic characters like soft hyphen (U+00AD) if they survived
                sanitized_label = sanitized_label.replace('\xad', '') # Remove U+00AD Soft Hyphen
                # Keep the label truncated to 40 chars for display clarity
                plotly_node_labels.append(sanitized_label[:40]) 
            except Exception as e:
                print(f"Sankey Diagram (v5) Warning: Error sanitizing label '{node_label[:50]}...': {e}. Using raw (may fail PNG).")
                plotly_node_labels.append(node_label[:40]) # Fallback to raw label
        else:
             # Handle non-string labels if they somehow occur
             plotly_node_labels.append(str(node_label)[:40])
        # --- END ADDED ---

        node_type = node_data.get("type") # e.g., "concept_item", "image_item", "task_box"
        agent_short_name = node_data.get("agent_short_name") # e.g., "ConceptGen", "ImgGen"

        color = default_node_color
        # Priority: Agent specific color, then type specific, then default
        if node_type == "task_box" and agent_short_name and agent_short_name in agent_color_map:
            color = agent_color_map[agent_short_name]
        elif node_type and node_type in color_map_by_type:
            color = color_map_by_type[node_type]
        elif agent_short_name and agent_short_name in agent_color_map: # If type is not task_box but agent info is there
            color = agent_color_map[agent_short_name]
        
        node_indices_map[node_id] = i
        plotly_node_colors.append(color)

    sources, targets, values, link_colors = [], [], [], []
    
    # --- MODIFIED: Default link color and conversion function ---
    default_link_hex = morandi_colors[16 % num_morandi] 
    default_link_alpha = 0.75 # Alpha as float for rgba

    def hex_to_rgba(hex_color, alpha):
        hex_color = hex_color.lstrip('#')
        r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        return f'rgba({r},{g},{b},{alpha})'

    default_link_rgba = hex_to_rgba(default_link_hex, default_link_alpha)
    # --- END MODIFIED ---

    for link_data in structured_links:
        source_id = link_data.get("source")
        target_id = link_data.get("target")
        
        # --- ADDED: More robust check for missing source/target ---
        if not source_id or not target_id:
            print(f"Sankey Diagram (v5) CRITICAL Warning: Link missing 'source' or 'target'. Link data: {link_data}. SKIPPING THIS LINK.")
            continue
        # --- END ADDED ---

        value = float(link_data.get("value", 1.0)) 

        if source_id not in node_indices_map:
            print(f"Sankey Diagram (v5) Warning: Source node ID '{source_id}' not found in mapped nodes. Skipping link from {source_id} to {target_id}.")
            continue
        if target_id not in node_indices_map:
            print(f"Sankey Diagram (v5) Warning: Target node ID '{target_id}' not found in mapped nodes. Skipping link from {source_id} to {target_id}.")
            continue

        sources.append(node_indices_map[source_id])
        targets.append(node_indices_map[target_id])
        values.append(max(0.1, value)) 
        
        link_color_from_llm_hex = link_data.get("color") # Assuming LLM might still provide hex
        link_alpha_from_llm = float(link_data.get("alpha", default_link_alpha)) # Allow LLM to specify alpha

        if link_color_from_llm_hex and isinstance(link_color_from_llm_hex, str) and link_color_from_llm_hex.startswith('#'):
            try:
                link_colors.append(hex_to_rgba(link_color_from_llm_hex, link_alpha_from_llm))
            except ValueError:
                print(f"Sankey Diagram (v5) Warning: Invalid hex color '{link_color_from_llm_hex}' from LLM for link {source_id}->{target_id}. Using default.")
                link_colors.append(default_link_rgba)
        elif link_color_from_llm_hex and isinstance(link_color_from_llm_hex, str) and link_color_from_llm_hex.startswith('rgba'): # if LLM provides full rgba
             link_colors.append(link_color_from_llm_hex)
        else:
            link_colors.append(default_link_rgba)


    if not sources:
        print("Sankey Diagram (v5): No valid links derived from structure data. Skipping plot.")
        return None

    fig = go.Figure(data=[go.Sankey(
        arrangement="snap", # CHANGED from "snap" to "perpendicular"
        node=dict(
            pad=25, # INCREASED padding slightly from 20 to 25
            thickness=15,
            line=dict(color="black", width=0.5),
            label=plotly_node_labels,
            color=plotly_node_colors,
        ),
        link=dict(
            source=sources, target=targets, value=values,
            color=link_colors,
            arrowlen=0 # --- MODIFIED: Remove arrow ---
        ))])
    
    fig.update_layout(
        # --- MODIFIED: Title and centering ---
        title_text="執行過程(參考)",  # UPDATED TITLE
        title_x=0.5, # Center title
        # --- END MODIFIED ---
        font_size=12,
        height=max(600, len(structured_nodes) * 25 + 100), # Added a bit more base height
        width=max(900, num_tasks_for_width_heuristic * 60 + 400), # Adjusted heuristic
        paper_bgcolor='rgba(255,255,255,1)', # White background
        plot_bgcolor='rgba(255,255,255,1)',  # White plot area
    )
    
    sankey_filename_base = f"workflow_sankey_v5_morandi_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    sankey_filepath_png = Path(output_dir) / f"{sankey_filename_base}.png"
    sankey_filepath_html = Path(output_dir) / f"{sankey_filename_base}.html"

    try:
        fig.write_image(str(sankey_filepath_png), scale=1.5)
        print(f"Sankey diagram (v5 Morandi) saved to: {sankey_filepath_png}")
        return str(sankey_filepath_png)
    except Exception as e_img:
        print(f"Error saving Sankey (v5 Morandi) as PNG: {e_img}")
        try:
            fig.write_html(str(sankey_filepath_html))
            print(f"Sankey diagram (v5 Morandi) saved as HTML (fallback): {sankey_filepath_html}")
            return str(sankey_filepath_html)
        except Exception as e_html:
            print(f"Error saving Sankey (v5 Morandi) as HTML: {e_html}")
            return None


def _add_detailed_evaluation_to_doc(
    doc: DocxDocument, 
    eval_task_data: TaskState, 
    base_heading_level: int,
    final_summary_output_dir: Path,
    current_task_id_for_report: str, # Used for asset path if not latest
    is_latest_eval_task: bool # To determine chart asset path
):
    """
    Adds a detailed evaluation section (summary, charts, scores table) to the Word document.
    """
    print(f"    Helper: Adding detailed evaluation for task {eval_task_data.get('task_id', 'N/A')}, base_level: {base_heading_level}")

    evaluation_content = eval_task_data.get("evaluation", {})
    detailed_scores_data = evaluation_content.get("detailed_assessment") or evaluation_content.get("detailed_option_scores")
    overall_assessment_text = evaluation_content.get("assessment", "N/A")
    overall_feedback_text = evaluation_content.get("final_llm_feedback_overall", evaluation_content.get("feedback_llm_overall", "N/A"))
    
    task_output_files = eval_task_data.get("output_files", [])
    charts_for_this_eval = []
    for f_info in task_output_files:
        if isinstance(f_info, dict):
            path_str = f_info.get("path")
            file_type = (f_info.get("type") or "").lower()
            name_str = f_info.get("filename") or (Path(path_str).name if path_str else "")
            desc_lower = (f_info.get("description") or "").lower()
            if path_str and Path(path_str).exists() and "image" in file_type:
                if "evaluation_radar" in name_str.lower() or \
                   "evaluation_stacked_bar" in name_str.lower() or \
                   "radar_chart" in desc_lower or \
                   "bar_chart" in desc_lower:
                    charts_for_this_eval.append(f_info)
                    print(f"    Helper: Found chart '{name_str}' for task {eval_task_data.get('task_id')}")

    if not detailed_scores_data and not charts_for_this_eval:
        doc.add_paragraph(f"（任務 {current_task_id_for_report}: 未找到詳細評估數據或相關圖表。）").italic = True
        doc.add_paragraph()
        # No page break here, let the caller handle page breaks between tasks.
        return

    # --- Overall Assessment and Feedback ---
    doc.add_heading("整體評估與回饋摘要", level=base_heading_level)
    p_assess = doc.add_paragraph(); p_assess.add_run("整體評估: ").bold = True; p_assess.add_run(overall_assessment_text)
    p_feed = doc.add_paragraph(); p_feed.add_run("整體 LLM 回饋: ").bold = True; p_feed.add_run(overall_feedback_text)
    doc.add_paragraph()

    # --- Evaluation Charts ---
    if charts_for_this_eval:
        doc.add_heading("評估圖表總覽", level=base_heading_level)
        doc.add_paragraph("以下圖表視覺化呈現了設計方案在各評估指標上的表現。")
        charts_for_this_eval.sort(key=lambda x: 0 if "evaluation_radar" in (x.get("filename") or "").lower() else 1)

        for chart_info in charts_for_this_eval:
            chart_path = chart_info.get("path")
            chart_name = chart_info.get("filename") or (Path(chart_path).name if chart_path else "N/A")
            chart_desc = chart_info.get("description") or chart_name

            if chart_path:
                p_ci1 = doc.add_paragraph(); p_ci1.add_run(f"File: SourceAgent: {chart_info.get('SourceAgent', 'N/A')}; TaskDesc: {chart_info.get('TaskDescShort', chart_desc)}; ImageNum: {chart_info.get('ImageNum', '')}").bold = True
                p_ci2 = doc.add_paragraph(); p_ci2.add_run("檔案位置: ").bold = True; p_ci2.add_run(chart_path)
                p_ci3 = doc.add_paragraph(); run_cd = p_ci3.add_run(f"(檔名: {chart_name}, 類型: Image/png)"); run_cd.italic = True; run_cd.font.size = Pt(9)
                doc.add_paragraph()

            if chart_path and Path(chart_path).exists():
                try:
                    if is_latest_eval_task:
                        asset_dir = final_summary_output_dir / "eval_charts"
                    else:
                        asset_dir = final_summary_output_dir / "task_assets" / current_task_id_for_report / "eval_charts"
                    os.makedirs(asset_dir, exist_ok=True)
                    dest_chart_file = asset_dir / Path(chart_path).name
                    shutil.copy2(chart_path, dest_chart_file)
                    
                    try:
                        p_img = doc.add_paragraph(); p_img.add_run().add_picture(str(dest_chart_file), width=Inches(6.0)); p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception: # PIL Fallback
                        img_pil = Image.open(str(dest_chart_file))
                        if img_pil.mode == 'RGBA': img_pil = img_pil.convert('RGB')
                        temp_pil_path = dest_chart_file.with_suffix(".pil_temp.jpg")
                        img_pil.save(temp_pil_path, "JPEG")
                        p_img_pil = doc.add_paragraph(); p_img_pil.add_run().add_picture(str(temp_pil_path), width=Inches(6.0)); p_img_pil.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if temp_pil_path.exists(): temp_pil_path.unlink()
                        chart_desc += " (PIL 處理)"
                    
                    cap_p = doc.add_paragraph(); cap_r = cap_p.add_run(f"圖: {chart_desc}"); cap_r.italic = True; cap_r.font.size = Pt(9); cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except Exception as e_chart_embed:
                    print(f"    Helper Error: Embedding chart '{chart_name}' failed: {e_chart_embed}")
                    doc.add_paragraph(f"  [嵌入圖表 '{chart_name}' 失敗: {e_chart_embed}]"); doc.add_paragraph()
            elif chart_path:
                doc.add_paragraph(f"  [圖表檔案未找到: {chart_name} (路徑: {chart_path})]"); doc.add_paragraph()
        doc.add_paragraph()

    # --- Detailed Scores Table ---
    if detailed_scores_data and isinstance(detailed_scores_data, list):
        if charts_for_this_eval: doc.add_page_break() # Page break if charts were present
        doc.add_heading("各方案詳細評估分數", level=base_heading_level)
        
        score_keys_display = [
            ("user_goal_responsiveness_score_final", "使用者目標響應度"),
            ("aesthetics_context_score_final", "美學與情境契合度"),
            ("functionality_flexibility_score_final", "功能性與彈性"),
            ("durability_maintainability_score_final", "耐久性與可維護性"),
            ("cost_efficiency_score_final", "成本效益"),
            ("green_building_score_final", "綠建築永續潛力")
        ]
        for option_data in detailed_scores_data:
            if isinstance(option_data, dict):
                opt_id = option_data.get("option_id", "未知方案")
                doc.add_heading(f"方案: {opt_id}", level=base_heading_level + 1)
                opt_desc = option_data.get("description", "無描述")
                p_opt_desc = doc.add_paragraph(); p_opt_desc.add_run("描述: ").bold = True; p_opt_desc.add_run(opt_desc)
                
                doc.add_heading("分數:", level=base_heading_level + 2)
                tbl = doc.add_table(rows=1, cols=2); tbl.style = 'Table Grid'
                hdr_cells = tbl.rows[0].cells; hdr_cells[0].text = '評估標準'; hdr_cells[1].text = '分數 (0-10)'
                for skey, dname in score_keys_display:
                    r_cells = tbl.add_row().cells; r_cells[0].text = dname
                    sval = option_data.get(skey)
                    r_cells[1].text = f"{sval:.1f}" if isinstance(sval, float) else (str(sval) if sval is not None else "N/A")
                
                rat = option_data.get("scoring_rationale", "無詳細理由")
                doc.add_heading("分數理由:", level=base_heading_level + 2); doc.add_paragraph(rat)
                doc.add_paragraph()
        # No page break after the last option table; outer loop will handle page breaks between tasks.
    
    # Add a final page break if this detailed section itself was substantial, 
    # to ensure separation before the next task in the "Workflow Details" or before "End of Report".
    # Only if content was actually added.
    if detailed_scores_data or charts_for_this_eval:
        if not (is_latest_eval_task and not detailed_scores_data): # Avoid double page break if latest eval has only charts
             if detailed_scores_data: # If table was added, always break.
                doc.add_page_break()


async def save_final_summary(state: WorkflowState, config: RunnableConfig) -> Dict[str, Any]:
    """
    Saves the final summary of the workflow...
    """
    from docx.shared import Inches, Pt # Keep local imports for clarity if helper is outside
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    print("--- Saving Final Workflow Summary (with AI Sankey Structure) ---")
    if not state:
        print("Error: State is None, cannot save summary.")
        # Return a dict as per type hint
        return {"error": "Input state was None", "current_phase": state.get("current_phase", "unknown")}


    final_summary_output_dir_str = "D:/MA system/LangGraph/output/Report" 
    final_summary_output_dir = Path(final_summary_output_dir_str) 
    os.makedirs(final_summary_output_dir_str, exist_ok=True)

    sankey_processed_state = state.copy() 

    if 'sankey_structure_agent' not in globals() or not isinstance(sankey_structure_agent, SankeyStructureAgent):
        print("FATAL ERROR: sankey_structure_agent not defined or not an instance of SankeyStructureAgent.")
        sankey_processed_state["sankey_structure_data"] = {"error": "SankeyStructureAgent not available."}
    else:
        try:
            sankey_processed_state_after_agent = await sankey_structure_agent.run(sankey_processed_state, config)
            sankey_processed_state = sankey_processed_state_after_agent
        except Exception as e_sankey_agent:
            print(f"Error running SankeyStructureAgent: {e_sankey_agent}")
            traceback.print_exc()
            sankey_processed_state["sankey_structure_data"] = {"error": f"SankeyStructureAgent execution failed: {e_sankey_agent}"}

    sankey_diagram_path = None
    # ... (Sankey diagram generation logic - unchanged) ...
    sankey_structure_for_plot = sankey_processed_state.get("sankey_structure_data")
    user_input_title = sankey_processed_state.get("user_input", "N/A")
    num_tasks_heuristic = len(sankey_processed_state.get("tasks", []))
    
    try:
        sankey_diagram_path = generate_sankey_diagram_plotly(
            sankey_structure_data=sankey_structure_for_plot,
            user_input_for_title=user_input_title,
            output_dir=final_summary_output_dir_str,
            num_tasks_for_width_heuristic=num_tasks_heuristic
        )
    except Exception as e_sankey_gen:
        print(f"Error during Sankey diagram generation call (v5): {e_sankey_gen}")
        traceback.print_exc()

    original_tasks_for_report_generation = sankey_processed_state.get("tasks", [])
    # ... (Debug prints for tasks - unchanged) ...

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    # ... (Filenames and doc initialization - unchanged) ...
    base_filename = f"Final_Workflow_Summary_{timestamp}"
    word_filename = f"{base_filename}.docx"
    json_filename = f"{base_filename}.json"
    word_filepath = os.path.join(final_summary_output_dir, word_filename)
    json_filepath = os.path.join(final_summary_output_dir, json_filename)
    doc = DocxDocument()
    
    # --- Report Title, Sankey Embedding, User Goal, Report Date --- (unchanged)
    title_paragraph = doc.add_heading('Workflow Final Summary', level=0)
    if title_paragraph.runs: title_paragraph.runs[0].bold = True
    doc.add_paragraph()

    if sankey_diagram_path and Path(sankey_diagram_path).suffix.lower() in ['.png', '.jpg', '.jpeg']:
        try:
            doc.add_paragraph("以下桑基圖展示了本次設計工作流程中主要的方案生成、分支探索、評估與收斂的路徑。")
            doc.add_picture(sankey_diagram_path, width=Inches(6.0))
            sankey_caption_para = doc.add_paragraph()
            sankey_caption_run = sankey_caption_para.add_run(f"圖1: 設計流程桑基圖總覽 ({Path(sankey_diagram_path).name})")
            sankey_caption_run.italic = True; sankey_caption_run.font.size = Pt(9)
            sankey_caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        except Exception as e_embed_sankey:
            print(f"Error embedding Sankey diagram into Word: {e_embed_sankey}")
            doc.add_paragraph(f"[無法自動嵌入桑基圖 '{Path(sankey_diagram_path).name}'. 錯誤: {e_embed_sankey}]"); doc.add_paragraph()
    elif sankey_diagram_path:
        doc.add_paragraph(f"桑基圖已生成，請參考檔案: {sankey_diagram_path}"); doc.add_paragraph()

    doc.add_paragraph(f"User Goal: {state.get('user_input', 'N/A')}") # Use original state for user_input
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(); doc.add_page_break()

    # --- Find latest_eval_task --- (unchanged)
    latest_eval_task = None
    for task_idx in range(len(original_tasks_for_report_generation) - 1, -1, -1):
        task_eval_check = original_tasks_for_report_generation[task_idx]
        if task_eval_check.get("selected_agent") in ["SpecialEvaAgent", "FinalEvaAgent"] and \
           task_eval_check.get("status") == "completed":
            latest_eval_task = task_eval_check; break
    print(f"Save Summary: Found latest relevant eval task (if any): ID {latest_eval_task.get('task_id') if latest_eval_task else 'N/A'}")

    # --- Workflow Details Section ---
    doc.add_heading("工作流程細節 (Workflow Details)", level=1)
    if not original_tasks_for_report_generation:
        doc.add_paragraph("No tasks were executed in this workflow.")
    else:
        for i, task in enumerate(original_tasks_for_report_generation):
            task_status = task.get("status", "N/A")
            task_agent = task.get("selected_agent")
            task_id_for_loop = task.get('task_id', f'task_idx_{i}')

            if latest_eval_task and task_id_for_loop == latest_eval_task.get('task_id'):
                continue # Skip latest_eval_task here

            if task_status in ["completed", "failed", "max_retries_reached"]:
                doc.add_heading(f"Task {i+1}: {task.get('description', 'N/A')}", level=2)
                # Basic Task Info (Objective, Agent, Status)
                p_obj = doc.add_paragraph(); p_obj.add_run("Objective: ").bold = True; p_obj.add_run(task.get('task_objective', 'N/A'))
                p_ag = doc.add_paragraph(); p_ag.add_run("Agent: ").bold = True; p_ag.add_run(task_agent or 'N/A')
                p_stat = doc.add_paragraph(); p_stat.add_run("Status: ").bold = True; p_stat.add_run(task_status)
                doc.add_paragraph()

                # Evaluation Criteria/Rubric
                if task.get("evaluation") and isinstance(task["evaluation"], dict):
                    # ... (Unchanged Evaluation Criteria/Rubric display logic) ...
                    specific_criteria = task["evaluation"].get("specific_criteria")
                    is_specific_criteria_valid_and_not_default = ( specific_criteria and isinstance(specific_criteria, str) and specific_criteria.strip() and specific_criteria.lower() not in ["default criteria apply / rubric not generated.","default criteria apply / rubric not generated","default criteria apply"])
                    if is_specific_criteria_valid_and_not_default or True: 
                        doc.add_heading("Evaluation Criteria/Rubric:", level=3)
                        if is_specific_criteria_valid_and_not_default:
                            doc.add_paragraph(specific_criteria); doc.add_paragraph(); doc.add_paragraph().add_run("通用評估標準補充說明：").bold = True
                        else: doc.add_paragraph("（以下為通用評估標準說明）").italic = True
                        p_cost_title = doc.add_paragraph(); p_cost_title.add_run("早期成本效益估算說明：").bold = True
                        doc.add_paragraph("有預算上限時：因為屬於前期成本概算，設定成本偏差閾值 ±50%計算得分。低於預算50%（即預算 * 0.5）因可能低於合理標的底價，視為1分；高於預算50%（即預算 * 1.5）因成本效益過低，亦視為1分。在預算 ±50%範圍內，成本越低（越接近預算 * 0.5），分數越高，呈線性關係.")
                        doc.add_paragraph("無預算上限時：基於成本效率分數計算，通常將觀察到的成本範圍（例如從最低成本到最高成本）進行線性映射給分，成本越低，分數越高.")
                        doc.add_paragraph() 
                        p_green_title = doc.add_paragraph(); p_green_title.add_run("綠建築永續潛力估算說明：").bold = True
                        doc.add_paragraph("大致基於綠建築標章之主要評估指標（如生態、健康、節能、減廢等四大項）進行計分。潛力分數的計算方式可能為各指標預期得分的加權總和，再轉換為0-10分制（例如，總達成率百分比除以10）。")
                        doc.add_paragraph("此潛力分數可對應至業界常見的綠建築評級潛力：3分以下約為合格級；3-6分約為銅級；6-8分約為銀級；8-9.5分約為黃金級；9.5分以上則具備鑽石級潛力。")
                        doc.add_paragraph()
                
                # --- MODIFIED: Call helper for intermediate detailed evaluation tasks ---
                is_intermediate_detailed_eval = task_agent in ["SpecialEvaAgent", "FinalEvaAgent"]
                
                if is_intermediate_detailed_eval:
                    print(f"Save Summary: Task {task_id_for_loop} is an intermediate eval. Calling helper.")
                    _add_detailed_evaluation_to_doc(
                        doc, task, 
                        base_heading_level=3, # Starts under "Task X" (level 2)
                        final_summary_output_dir=final_summary_output_dir,
                        current_task_id_for_report=task_id_for_loop,
                        is_latest_eval_task=False
                    )
                    # The helper function now handles charts and detailed scores.
                    # The generic "Textual Outputs" and "Associated Files" below should be skipped for these tasks.
                else:
                    # --- Standard Textual Outputs and Associated Files for NON-DETAILED tasks ---
                    doc_text_heading_added = False # Reset for each task
                    # ... (Existing general textual output logic - unchanged) ...
                    # Part 1: Specific Evaluation Summary (This part might be redundant if covered by helper for eval tasks)
                    # However, the condition `task_agent in ["SpecialEvaAgent", "FinalEvaAgent"]` is now handled by `is_intermediate_detailed_eval`
                    # So this block will effectively not run if `is_intermediate_detailed_eval` is true.
                    # If it's a *non-detailed* eval agent (if such a thing exists), this might still apply.
                    # For clarity, let's keep the condition specific.
                    # This section will ONLY run if `is_intermediate_detailed_eval` is FALSE.
                    
                    # The previous code had a section here:
                    # if task_agent in ["SpecialEvaAgent", "FinalEvaAgent"]: (and it's not latest_eval_task)
                    # This is now covered by the `is_intermediate_detailed_eval` and the call to the helper.
                    # So, we only need the "General Text Outputs" part here for non-detailed tasks.

                    # Part 2: General Text Outputs from task["outputs"]
                    general_outputs_list = []
                    if task.get("outputs"):
                        excluded_gen_keys = ["mcp_internal_messages", "grounding_sources", "search_suggestions", "radar_chart_path", "detailed_option_scores", "detailed_assessment", "assessment", "feedback_llm_overall", "final_llm_feedback_overall"]
                        for k, v_str in task["outputs"].items():
                            if isinstance(v_str, str) and k not in excluded_gen_keys and len(v_str) > 10: general_outputs_list.append((k, v_str))
                    if general_outputs_list:
                        if not doc_text_heading_added: doc.add_heading("關鍵文字輸出 (Key Text Outputs):", level=3); doc_text_heading_added = True
                        for k_out, v_out in general_outputs_list: doc.add_paragraph(f"{k_out.replace('_', ' ').title()}:", style='Intense Quote'); doc.add_paragraph(v_out); doc.add_paragraph()
                    if doc_text_heading_added or general_outputs_list: doc.add_paragraph()


                    # Associated Files for NON-DETAILED tasks
                    if task.get("output_files"):
                        doc.add_heading("關聯檔案 (Associated Files):", level=3)
                        for file_info_gen in task.get("output_files", []):
                            if not isinstance(file_info_gen, dict): continue
                            path_gen_str = file_info_gen.get("path")
                            type_gen = (file_info_gen.get("type") or "").lower()
                            name_gen = file_info_gen.get("filename") or (Path(path_gen_str).name if path_gen_str else "N/A")
                            desc_gen = file_info_gen.get("description") or name_gen

                            is_eval_chart_generic = False # Check if it's an eval chart to skip
                            if path_gen_str and Path(path_gen_str).exists() and "image" in type_gen:
                                if "evaluation_radar" in name_gen.lower() or "evaluation_stacked_bar" in name_gen.lower() or \
                                   "radar_chart" in desc_gen.lower() or "bar_chart" in desc_gen.lower():
                                    is_eval_chart_generic = True
                            
                            if is_eval_chart_generic: # Skip eval charts here, handled by helper
                                print(f"Save Summary (Workflow Details): Skipping eval chart {name_gen} in generic files, handled by detailed section.")
                                continue
                            
                            p_f_info1 = doc.add_paragraph(); p_f_info1.add_run(f"File: SourceAgent: {file_info_gen.get('SourceAgent', 'N/A')}; TaskDesc: {file_info_gen.get('TaskDescShort', desc_gen)}; ImageNum: {file_info_gen.get('ImageNum', '')}").bold = True
                            p_f_info2 = doc.add_paragraph(); p_f_info2.add_run("檔案位置: ").bold = True; p_f_info2.add_run(path_gen_str or "N/A")
                            p_f_info3 = doc.add_paragraph(); run_f_details = p_f_info3.add_run(f"(檔名: {name_gen or 'N/A'}, 類型: {type_gen.capitalize() or 'Unknown'})"); run_f_details.italic = True; run_f_details.font.size = Pt(9)
                            doc.add_paragraph()
                            if path_gen_str and Path(path_gen_str).exists():
                                asset_dir_gen = Path(final_summary_output_dir) / "task_assets" / task_id_for_loop
                                os.makedirs(asset_dir_gen, exist_ok=True)
                                try:
                                    dest_file_gen = asset_dir_gen / Path(path_gen_str).name
                                    shutil.copy2(path_gen_str, dest_file_gen)
                                    if "image" in type_gen:
                                        try:
                                            p_gen_img = doc.add_paragraph(); p_gen_img.add_run().add_picture(str(dest_file_gen), width=Inches(5.0)); p_gen_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            cap_p_gen = doc.add_paragraph(); cap_r_gen = cap_p_gen.add_run(f"圖: {desc_gen}"); cap_r_gen.italic = True; cap_r_gen.font.size = Pt(9); cap_p_gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            doc.add_paragraph()
                                        except Exception: 
                                            img_pil_gen = Image.open(str(dest_file_gen))
                                            if img_pil_gen.mode == 'RGBA': img_pil_gen = img_pil_gen.convert("RGB")
                                            temp_pil_path_gen = dest_file_gen.with_suffix(".pil_temp.jpg")
                                            img_pil_gen.save(temp_pil_path_gen, "JPEG")
                                            p_gen_img_pil = doc.add_paragraph(); p_gen_img_pil.add_run().add_picture(str(temp_pil_path_gen), width=Inches(5.0)); p_gen_img_pil.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            cap_p_gen_pil = doc.add_paragraph(); cap_r_gen_pil = cap_p_gen_pil.add_run(f"圖: {desc_gen} (PIL 處理)"); cap_r_gen_pil.italic = True; cap_r_gen_pil.font.size = Pt(9); cap_p_gen_pil.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            doc.add_paragraph()
                                            if temp_pil_path_gen.exists(): temp_pil_path_gen.unlink()
                                    elif "video" in type_gen:
                                        print(f"Save Summary: File {name_gen} is a video. Attempting to extract frames at 12s and 18s for generic task.")
                                        video_path_for_frames = str(dest_file_gen) # Use the copied video path
                                        temp_frame_paths_for_cleanup = [] 
                                        # Create a temporary sub-directory within the task's asset folder for frames
                                        frames_temp_dir = asset_dir_gen / "_frames_temp"
                                        frames_temp_dir.mkdir(parents=True, exist_ok=True)

                                        try:
                                            vidcap = cv2.VideoCapture(video_path_for_frames)
                                            if not vidcap.isOpened():
                                                print(f"Save Summary Warning: Could not open video file {video_path_for_frames}")
                                                doc.add_paragraph(f"  [無法打開影片檔案 '{name_gen}']")
                                            else:
                                                fps = vidcap.get(cv2.CAP_PROP_FPS)
                                                frame_count = int(vidcap.get(cv2.CAP_PROP_FRAME_COUNT))
                                                duration = frame_count / fps if fps > 0 else 0
                                                print(f"Save Summary Debug: Video {name_gen} - FPS: {fps}, Frame Count: {frame_count}, Duration: {duration:.2f}s")
                                                timestamps_seconds = [12, 18]
                                                for ts in timestamps_seconds:
                                                    if duration > ts:
                                                        target_frame_number = int(ts * fps)
                                                        vidcap.set(cv2.CAP_PROP_POS_FRAMES, target_frame_number)
                                                        success, image = vidcap.read()
                                                        if success:
                                                            frame_filename_ts = f"{Path(name_gen).stem}_{ts}s.jpg"
                                                            temp_frame_path = frames_temp_dir / frame_filename_ts
                                                            cv2.imwrite(str(temp_frame_path), image)
                                                            temp_frame_paths_for_cleanup.append(temp_frame_path)
                                                            try:
                                                                doc.add_picture(str(temp_frame_path), width=Inches(5.0))
                                                                para_caption_frame = doc.add_paragraph()
                                                                run_caption_frame = para_caption_frame.add_run(f"圖: 影片 '{name_gen}' 在 {ts} 秒的畫面")
                                                                run_caption_frame.italic = True; run_caption_frame.font.size = Pt(9)
                                                                para_caption_frame.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                                doc.add_paragraph()
                                                            except Exception as embed_frame_err:
                                                                 print(f"Save Summary Error: Failed to embed frame for {name_gen} at {ts}s: {embed_frame_err}")
                                                                 doc.add_paragraph(f"  [無法自動嵌入影片 '{name_gen}' 在 {ts} 秒的畫面。錯誤: {embed_frame_err}]")
                                                        else:
                                                            print(f"Save Summary Warning: Could not read frame at {ts}s from video {name_gen}.")
                                                            doc.add_paragraph(f"  [無法讀取影片 '{name_gen}' 在 {ts} 秒的畫面]")
                                                    else:
                                                         print(f"Save Summary Warning: Video {name_gen} is shorter than {ts}s. Skipping frame at {ts}s.")
                                                         doc.add_paragraph(f"  [影片 '{name_gen}' 時長不足 {ts} 秒，無法擷取該時間點畫面]")
                                            vidcap.release()
                                        except Exception as frame_extract_err:
                                            print(f"Save Summary Error: Error extracting/embedding frames from video {name_gen}: {frame_extract_err}")
                                            doc.add_paragraph(f"  [處理影片 '{name_gen}' 時發生錯誤：{frame_extract_err}]")
                                        finally:
                                            for temp_path in temp_frame_paths_for_cleanup:
                                                if temp_path.exists():
                                                    try: temp_path.unlink()
                                                    except Exception as e_cleanup: print(f"Save Summary Warning: Failed to delete temp frame {temp_path}: {e_cleanup}")
                                            if frames_temp_dir.exists() and not any(frames_temp_dir.iterdir()): # Remove temp dir if empty
                                                try: frames_temp_dir.rmdir()
                                                except Exception as e_rmdir: print(f"Save Summary Warning: Failed to remove empty frames_temp_dir {frames_temp_dir}: {e_rmdir}")

                                        doc.add_paragraph(f"  [影片檔案 '{name_gen}' 已複製 - 請至 task_assets 資料夾觀看]"); doc.add_paragraph()
                                        
                                    elif "model" in type_gen or Path(path_gen_str).suffix.lower() in ['.glb', '.obj', '.fbx', '.stl']:
                                        doc.add_paragraph(f"  [3D 模型檔案 '{name_gen}' 已複製 - 請至 task_assets 資料夾觀看]"); doc.add_paragraph()
                                    else: doc.add_paragraph(f"  [檔案類型 '{type_gen}' ({name_gen}) 已複製，路徑已提供]"); doc.add_paragraph()
                                except Exception as copy_e_gen: doc.add_paragraph(f"\n  - 原始路徑: {path_gen_str} (複製錯誤: {copy_e_gen})"); doc.add_paragraph()
                            elif path_gen_str: doc.add_paragraph(f"\n  - 原始路徑 (檔案未找到): {path_gen_str}"); doc.add_paragraph()
                            else: doc.add_paragraph("\n  - 檔案資訊缺少 'path' 欄位。"); doc.add_paragraph()

                # Page break logic for "Workflow Details" section
                # (Unchanged page break logic between tasks in this section)
                has_subsequent_printable_task = False
                for j in range(i + 1, len(original_tasks_for_report_generation)):
                    next_task_in_list = original_tasks_for_report_generation[j]
                    is_next_task_latest_eval = latest_eval_task and next_task_in_list.get('task_id') == latest_eval_task.get('task_id')
                    if next_task_in_list.get("status") in ["completed", "failed", "max_retries_reached"] and not is_next_task_latest_eval:
                        has_subsequent_printable_task = True; break
                if has_subsequent_printable_task: doc.add_page_break()
    # --- End of Workflow Details loop ---

    # --- Detailed Assessment Results Section (for latest_eval_task) ---
    doc.add_heading("詳細評估結果 (Detailed Assessment Results)", level=1)
    if latest_eval_task:
        print(f"Save Summary: Calling helper for latest_eval_task ID: {latest_eval_task.get('task_id')}")
        _add_detailed_evaluation_to_doc(
            doc, latest_eval_task,
            base_heading_level=2,
            final_summary_output_dir=final_summary_output_dir,
            current_task_id_for_report=latest_eval_task.get('task_id', 'latest_eval'), # Should have ID
            is_latest_eval_task=True
        )
    else:
        doc.add_paragraph("未在已完成的評估任務中找到可供最終詳細評估的資料。")
        doc.add_paragraph()
        doc.add_page_break() # Add page break if this section is empty

    # --- End of Report, Save Doc, Save JSON --- (unchanged)
    doc.add_paragraph(); doc.add_heading("報告結束 (End of Report)", level=1)
    try:
        doc.save(word_filepath)
        print(f"Final summary Word document saved to: {word_filepath}")
    except Exception as e_save_doc:
        print(f"Error saving Word document: {e_save_doc}"); traceback.print_exc()
        try: fallback_word_path = final_summary_output_dir / f"Fallback_Summary_{timestamp}.docx"; doc.save(fallback_word_path); print(f"Saved Word document with fallback name: {fallback_word_path}")
        except Exception as fe_save_doc: print(f"Failed to save Word document with fallback name: {fe_save_doc}")

    try:
        serializable_state_for_file = {} 
        for key, value in sankey_processed_state.items(): 
            if key == "config": serializable_state_for_file[key] = "Configuration object (not fully serialized)"
            elif isinstance(value, Path): serializable_state_for_file[key] = str(value)
            elif key == "tasks" and isinstance(value, list): 
                 serializable_state_for_file[key] = []
                 for task_item_json in value: 
                     if isinstance(task_item_json, dict):
                         s_task_json = {} 
                         for t_key_json, t_val_json in task_item_json.items(): 
                             if t_key_json == "output_files" and isinstance(t_val_json, list):
                                 s_task_json[t_key_json] = []
                                 for f_item_json in t_val_json: 
                                     if isinstance(f_item_json, dict):
                                         f_copy_json = f_item_json.copy(); f_copy_json.pop("base64_data", None); s_task_json[t_key_json].append(f_copy_json)
                                     else: s_task_json[t_key_json].append(f_item_json)
                             elif isinstance(t_val_json, (dict, list, str, int, float, bool, type(None))): s_task_json[t_key_json] = t_val_json
                             else: s_task_json[t_key_json] = f"<Non-serializable type: {type(t_val_json).__name__}>"
                         serializable_state_for_file[key].append(s_task_json)
                     else: serializable_state_for_file[key].append(f"<Non-dict task item: {type(task_item_json).__name__}>")
            elif isinstance(value, (dict, list, str, int, float, bool, type(None))): serializable_state_for_file[key] = value
            else: serializable_state_for_file[key] = f"<Non-serializable type: {type(value).__name__}>"
        with open(json_filepath, 'w', encoding='utf-8') as f_json: # Distinct file handle
            json.dump(serializable_state_for_file, f_json, indent=4, ensure_ascii=False, default=str)
        print(f"Final summary JSON state saved to: {json_filepath}")
    except Exception as e_save_json:
        print(f"Error saving JSON state: {e_save_json}"); traceback.print_exc()
        try:
            fallback_json_path = final_summary_output_dir / f"Fallback_State_{timestamp}.json"
            with open(fallback_json_path, 'w', encoding='utf-8') as f_fallback_json:
                 json.dump({"error": "Failed to serialize full state, this is a minimal fallback.", "original_error": str(e_save_json)}, f_fallback_json, indent=4, ensure_ascii=False)
            print(f"Saved JSON state with fallback name: {fallback_json_path}")
        except Exception as fe_save_json: print(f"Failed to save JSON state with fallback name: {fe_save_json}")

    return_payload = {
        "final_summary_word_path": str(word_filepath), 
        "current_phase": "qa", 
        # "sankey_structure_data": sankey_processed_state.get("sankey_structure_data") # No cleaning
    }
    print(f"Final summary saved. Set current_phase to 'qa'. Returning payload to graph.")
    return return_payload

def initialize_workflow(user_input: str) -> WorkflowState:
    """Initializes the workflow state for a new run."""
    print(f"Initializing workflow for input: {user_input[:100]}...")
    initialized_state = WorkflowState(
        user_input=user_input,
        tasks=[],
        current_task_index=0,
        current_phase="task_execution",
        qa_context=[],
        interrupt_input=None,
        current_task=None,
        sankey_structure_data=None # Initialize new field
    )
    print(f"Initialized workflow state: Phase='{initialized_state['current_phase']}', user_input='{user_input[:50]}...'")
    return initialized_state



# =============================================================================
# 9. 圖邊緣條件函數定義 (Update routing)
# =============================================================================
def route_from_process_management(state: WorkflowState) -> Literal[
    "assign_teams", "eva_teams", "save_final_summary", "qa_loop", "finished", "process_management"
]:
    """ Determines the next node based on task status and selected agent."""
    print("--- Routing decision @ route_from_process_management (v2) ---") # Version indicator
    current_phase = state.get("current_phase")
    interrupt_result = state.get("interrupt_result")
    interrupt_action = interrupt_result.get("action") if isinstance(interrupt_result, dict) else None

    print(f"Router received state: Phase='{current_phase}', Interrupt Action='{interrupt_action}'")

    # --- 1. Check QA / Finished ---
    if current_phase == "qa" or interrupt_action == "CONVERSATION":
         state["interrupt_result"] = None
         print(f"Routing -> qa_loop (Reason: Phase='{current_phase}' or Interrupt Action='CONVERSATION')")
         return "qa_loop"
    if current_phase == "finished":
        print("Routing -> finished (Reason: Phase is 'finished')")
        return "finished"

    # --- 2. Normal Task Execution Flow ---
    tasks = state.get("tasks", [])
    current_index = state.get("current_task_index", 0)

    if not tasks or current_index >= len(tasks):
        if current_phase != "finished":
             print(f"Routing -> save_final_summary (Reason: All tasks completed or list empty, index={current_index}, count={len(tasks)})")
             return "save_final_summary"
        else:
             print("Routing -> finished (Reason: Tasks done and phase 'finished')")
             return "finished"


    if 0 <= current_index < len(tasks):
        current_task = tasks[current_index]
        task_id = current_task.get('task_id', 'N/A')
        task_status = current_task.get("status")
        selected_agent = current_task.get("selected_agent")
        requires_eval = current_task.get("requires_evaluation", False)

        print(f"  Checking Task {task_id} (Idx: {current_index}), Status: '{task_status}', Agent: '{selected_agent}', RequiresEval: {requires_eval}")

        if task_status in ["failed", "max_retries_reached"]:
            print(f"Routing -> process_management (Reason: Task {task_id} status '{task_status}', returning for failure analysis)")
            return "process_management"

        # --- MODIFIED: Check for specific evaluation agents ---
        is_evaluation_agent = selected_agent in ["EvaAgent", "SpecialEvaAgent", "FinalEvaAgent"]

        if is_evaluation_agent:
            if requires_eval:
                if task_status in ["pending", "in_progress", None]: # Should be pending
                    print(f"Routing -> eva_teams (Reason: Task {task_id} is agent '{selected_agent}', requires evaluation, and status '{task_status}')")
                    # --- CORRECTION: The target node is still called "eva_teams" which triggers EvaAgent.run ---
                    return "eva_teams" # Route to the single evaluation entry point
                else:
                    print(f"Routing Warning: Evaluation Task {task_id} ({selected_agent}) has unexpected status '{task_status}'. Looping PM.")
                    return "process_management"
            else:
                 print(f"Routing Error: Evaluation Task {task_id} ({selected_agent}) has requires_evaluation=False. Looping PM.")
                 # This indicates an error in PM's plan generation
                 return "process_management"
        else: # Not an evaluation agent
            if task_status in ["pending", "in_progress", None]:
                print(f"Routing -> assign_teams (Reason: Task {task_id} agent '{selected_agent}' status '{task_status}' and ready for execution)")
                return "assign_teams"
            else: # Should already be completed or failed, handled above
                print(f"Routing Warning: Non-Eval Task {task_id} has unexpected status '{task_status}'. Looping PM.")
                return "process_management"
    else:
        print(f"Routing Warning: index ({current_index}) out of bounds (size {len(tasks)}). Fallback to save_final_summary.")
        return "save_final_summary"


# ... (route_after_assign_teams保持不變) ...
def route_after_assign_teams(state: WorkflowState) -> Literal["process_management", "finished"]:
    """
    Routes after the assign_teams subgraph. ALWAYS routes back to Process Management
    for status checking, LTM saving, and potential evaluation triggering.
    """
    # --- <<< 增加詳細日誌 >>> ---
    print(f"--- Debug log @ route_after_assign_teams ---")
    tasks = state.get("tasks", [])
    current_idx = state.get("current_task_index", -1)
    print(f"  Received current_task_index: {current_idx}")
    if 0 <= current_idx < len(tasks):
        task_in_state = tasks[current_idx]
        status_in_state = task_in_state.get('status')
        task_id_in_state = task_in_state.get('task_id')
        print(f"  Task {task_id_in_state} at index {current_idx} has status: '{status_in_state}' in the received state.")
    else:
        print(f"  Invalid current_task_index ({current_idx}) or empty tasks list.")
    # --- <<< 結束增加日誌 >>> ---

    # Basic index check
    if current_idx < 0 or current_idx >= len(tasks):
        print(f"Routing Error: Invalid task index {current_idx} after assign_teams. Routing to finished.")
        return "finished"

    print(f"Route after assign_teams: Routing decision -> process_management")
    return "process_management"

# def route_after_evaluation(state: WorkflowState) -> Literal["process_management", "assign_teams", "finished"]:
#     """Routes after EvaAgent (evaluation subgraph): Continue (PM) or retry (assign_teams)."""
#     current_phase = state.get("current_phase")
#     if current_phase != "task_execution":
#         print(f"Route after Eva Error: Unexpected phase '{current_phase}'. Routing to finished.")
#         return "finished"

#     tasks = state.get("tasks", [])
#     current_idx = state.get("current_task_index", -1)
#     if current_idx < 0 or current_idx >= len(tasks):
#         print("Routing Error: Invalid index after evaluation. Routing to finished.")
#         return "finished"

#     task = tasks[current_idx]
#     status = task.get("status")
#     retry_count = task.get("retry_count", 0)
#     max_retries = _full_static_config.agents.get("process_management", {}).parameters.get("max_retries", 3)

#     print(f"Route after Eva: Checking Task {current_idx}, Status: {status}, RetryCount: {retry_count}")

#     if status == "completed":
#         print("Route after Eva: Evaluation passed. Routing to PM to advance.")
#         return "process_management"
#     elif status == "failed":
#         print(f"Route after Eva: Evaluation failed. Routing to PM for assessment.")
#         return "process_management"
#     else:
#         print(f"Route after Eva Error: Unexpected status '{status}'. Routing to PM as failsafe.")
#         return "process_management"


# --- <<< 修改 route_after_chat_bot >>> ---
def route_after_chat_bot(state: WorkflowState) -> Literal["process_management", "qa_loop", "finished"]:
    """基於QA代理設置的階段路由到適當的節點"""
    current_phase = state.get("current_phase")
    print(f"--- 路由決策 @ route_after_chat_bot ---")
    print(f"  當前階段: '{current_phase}'")
    
    # 直接根據QA代理設置的階段進行路由，不再進行重複檢測
    if current_phase == "task_execution":
        print("路由 -> process_management (原因: QA代理設置階段為'task_execution')")
        return "process_management"
    elif current_phase == "finished":
        print("路由 -> finished (原因: QA代理設置階段為'finished')")
        return "finished"
    elif current_phase == "qa":
        print("路由 -> qa_loop (原因: QA代理設置階段為'qa')")
        return "qa_loop"
    else:
        print(f"路由警告: 未預期的階段 '{current_phase}'。預設路由到qa_loop。")
        return "qa_loop"


def qa_loop_node(state: WorkflowState) -> Dict[str, Any]:
    """
    Manages the user input part of the QA loop.
    This node now primarily serves as an interrupt point.
    It does not modify interrupt_input or qa_context directly.
    QA_Agent will read interrupt_input.
    """
    node_name = "QA Loop Manager (User Input Point)"
    print(f"--- Running Node: {node_name} ---")
    
    # LangGraph將處理來自用戶界面/API調用的 interrupt_input 填充。
    # 此節點本身不再需要修改 interrupt_input 或 qa_context。
    # 它只是一個流程中的點，允許 LangGraph 在此處暫停以等待 interrupt_input。
    current_interrupt = state.get("interrupt_input")
    if current_interrupt:
        print(f"{node_name}: Detected interrupt_input: '{str(current_interrupt)[:100]}...'. Passing to QA_Agent.")
    else:
        print(f"{node_name}: No interrupt_input detected for this cycle. QA_Agent will handle.")
        
    return {} # 返回空字典，LangGraph 會合併狀態，interrupt_input（如果被用戶設置）會保留


# =============================================================================
# 10. 圖定義與節點/邊緣添加
# =============================================================================
workflow = StateGraph(WorkflowState, config_schema=ConfigSchema)

# Add Nodes (No changes here)
workflow.add_node("process_management", process_management.run)
workflow.add_node("assign_teams", assign_teams)
workflow.add_node("eva_teams", eva_agent.run) # Still points to the EvaAgent.run trigger
workflow.add_node("chat_bot", qa_agent.run)
workflow.add_node("save_final_summary", save_final_summary)
workflow.add_node("user_message", qa_loop_node)

# Set Entry Point
workflow.set_entry_point("process_management")

# Add Edges
workflow.add_conditional_edges(
    "process_management",
    route_from_process_management, # Uses updated logic
    {
        "assign_teams": "assign_teams",
        "eva_teams": "eva_teams", # Still route to the eva_teams node trigger
        "save_final_summary": "save_final_summary",
        "qa_loop": "user_message",
        "process_management": "process_management", # Loop back for failure/retries
        "finished": END,
    }
)

workflow.add_conditional_edges(
    "assign_teams",
    route_after_assign_teams, # Still goes back to PM
    {
        "process_management": "process_management",
        # "finished": END # Should ideally not happen if PM handles empty tasks
    }
)

# --- REMOVED Conditional edge for eva_teams ---
# The main router now handles the state after evaluation completes.
# workflow.add_conditional_edges(
#     "eva_teams",
#     route_after_evaluation, # REMOVED
#     {       
#        "process_management": "process_management",
#        "finished": END
# )
# --- Instead, eva_teams implicitly returns to the main router ---
# The output of eva_teams (which is the state updated by the subgraph)
# flows back to the main graph logic implicitly after the node finishes.
# The next routing decision happens at the START of the next cycle, triggered
# by the process_management node again reading the updated state.
# THEREFORE, we need an edge FROM eva_teams BACK to process_management.
workflow.add_edge("eva_teams", "process_management")


# QA Loop Edges (Remain same)
workflow.add_edge("user_message", "chat_bot")
workflow.add_conditional_edges(
    "chat_bot",
    route_after_chat_bot,
    {
        "process_management": "process_management",
        "qa_loop": "user_message",
        "finished": END
    }
)

# Final Summary Edge (Remains same)
workflow.add_edge("save_final_summary", "process_management")

# =============================================================================
# 11. 圖編譯
# =============================================================================
# Need to adjust interrupt points if EvaAgent.run is just a trigger
# Interrupting before eva_teams might be less useful now.
# Interrupting *inside* the subgraph might be better if needed.
# For now, keep existing interrupts.
graph = workflow.compile(interrupt_before=["user_message"], interrupt_after=["chat_bot"]).with_config({"recursion_limit": 1000})
# graph = workflow.compile() #沒有中斷點的編譯
graph.name = "General_Arch_graph_v21_AISankey" # Updated name
print("Main Graph compiled successfully (v21 - AI Sankey Structure).")

# =============================================================================
# 12. TODOLIST
# =============================================================================
# 1. QA test   V
# 2. from insert enter QA node test    V
# 3. Fauilure analysis test   V
# 4. Building life cycle tools add in  (maybe 3D model)  V
# 5. final configuration systematically...    V
# 6. 3D model input / output test(MCP)  V
# 7. Change node name to agent name    V
# 8. final evaluation 好像不太行 
# 9. 檢查save summary 與長期記憶有關   V

